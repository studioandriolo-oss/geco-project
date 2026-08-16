import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import graphviz
import streamlit.components.v1 as components
from datetime import datetime, date
import json
from io import BytesIO
from docx import Document
import html
import base64

# --- CONFIGURAZIONE PAGINA ---
st.set_page_config(page_title="WBS/OBS Manager & EVM", layout="wide", initial_sidebar_state="expanded")

# --- RIMOZIONE TOTALE HEADER, SIDEBAR E STILE COLONNA COMPATTA ---
st.markdown("""
<style>
header[data-testid="stHeader"] {display: none !important;}
#MainMenu, footer, [data-testid="stSidebar"], [data-testid="collapsedControl"] {display: none !important;}
.block-container {
    padding-top: 1rem !important;
    padding-bottom: 1rem !important;
    max-width: 98% !important;
}
.btn-compatto button {
    padding: 0.2rem 0.5rem !important;
    min-height: 30px !important;
    font-size: 0.85rem !important;
    margin-bottom: 0px !important;
}
</style>
""", unsafe_allow_html=True)

# ==========================================
# IL MOTORE (PRIMA DELLA GRAFICA)
# ==========================================

# --- SISTEMA DI LOGIN SICURO ---
try:
    USER_ID = st.secrets["USER_ID"]
    PASSWORD = st.secrets["PASSWORD"]
except KeyError:
    st.error("⚠️ Errore di sistema: Credenziali non trovate. Configura i 'Secrets' di Streamlit.")
    st.stop()

if st.query_params.get("auth") == "valid":
    st.session_state.logged_in = True
elif 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.markdown("<br><br><h2 style='text-align: center;'>🔒 Accesso Riservato GECO</h2>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        with st.form("login_form"):
            user_input = st.text_input("ID Utente")
            pass_input = st.text_input("Password", type="password")
            submit = st.form_submit_button("Accedi al Gestionale", use_container_width=True)
            
            if submit:
                if user_input == USER_ID and pass_input == PASSWORD:
                    st.session_state.logged_in = True
                    st.query_params["auth"] = "valid" 
                    st.rerun() 
                else:
                    st.error("Credenziali errate. Riprova.")                
    st.stop()

# --- 1. INIZIALIZZAZIONE DATI ---
if 'wbs_data' not in st.session_state:
    st.session_state.wbs_data = pd.DataFrame([{
        'ID_WBS': '1', 
        'Attività': 'Progetto Principale', 
        'Inizio_Previsto': None, 'Fine_Prevista': None, 
        'Inizio_Effettivo': None, 'Fine_Effettiva': None, 
        'BAC_Budget': 0.0,
        '%_Completamento': 0.0, 
        'AC_Costo_Reale': 0.0, 
        'ID_OBS_Assegnato': None, 
        'Predecessori': '',
        'Vincolo_Burocratico': 'nessuno',
        'Vincolo_Assolto': False
    }])
    
if 'obs_data' not in st.session_state:
    st.session_state.obs_data = pd.DataFrame(columns=[
        'ID_OBS', 'Ruolo', 'Risorsa', 'Tipo_Contratto', 'Note'
    ])
    
if 'registro_data' not in st.session_state:
    st.session_state.registro_data = pd.DataFrame(columns=[
        'Data', 'N_Doc', 'Fornitore', 'Voce_WBS', 'Importo_Netto', 'Descrizione'
    ])
    
if 'tickets_data' not in st.session_state:
    st.session_state.tickets_data = pd.DataFrame(columns=[
        'ID_Ticket', 'ID_WBS_Rif', 'Autore', 'Data_Apertura',
        'Tipologia', 'Descrizione', 'Stato', 'Risposta_RUP', 'Data_Chiusura', 'Variazione_Costi', 'Variazione_Tempi', 'Variante_Applicata', 
    ])

if 'capa_data' not in st.session_state:
    st.session_state.capa_data = pd.DataFrame(columns=[
        'Data_Apertura', 'ID_WBS_Rif', 'Tipo_Azione', 'Descrizione', 'Responsabile_OBS', 'Stato', 'Rischio_Associato', 'Costo_Intervento', 'Giorni_Intervento', 'Costo_Scaricato'
    ])

if 'rischi_data' not in st.session_state:
    st.session_state.rischi_data = pd.DataFrame(columns=[
        'ID_WBS_Rif', 'Descrizione_Rischio', 'Probabilità (1-5)', 'Impatto (1-5)', 'Stato'
    ])

if 'sal_data' not in st.session_state:
    st.session_state.sal_data = pd.DataFrame(columns=['Data_Emissione', 'Descrizione_SAL', 'Importo_Euro', 'Stato_Pagamento'])

if 'archivio_progetti' not in st.session_state:
    st.session_state.archivio_progetti = {}
if 'nome_progetto_attivo' not in st.session_state:
    st.session_state.nome_progetto_attivo = "Nuovo_Progetto"


# --- 2. MOTORI MATEMATICI ---
def aggiorna_gerarchia(df):
    df_calc = df.copy()
    df_calc['BAC_Budget'] = pd.to_numeric(df_calc['BAC_Budget'], errors='coerce').fillna(0.0)
    df_calc['AC_Costo_Reale'] = pd.to_numeric(df_calc['AC_Costo_Reale'], errors='coerce').fillna(0.0)
    df_calc['%_Completamento'] = pd.to_numeric(df_calc['%_Completamento'], errors='coerce').fillna(0.0)
    
    for col in ['Inizio_Previsto', 'Fine_Prevista', 'Inizio_Effettivo', 'Fine_Effettiva']:
        if col in df_calc.columns:
            df_calc[col] = df_calc[col].astype(object)
            
    ids = df_calc['ID_WBS'].astype(str).tolist()
    foglie = [uid for uid in ids if not any(other.startswith(uid + '.') for other in ids if other != uid)]
    df_calc['Is_Leaf'] = df_calc['ID_WBS'].astype(str).isin(foglie)
    
    df_calc['Livello'] = df_calc['ID_WBS'].astype(str).apply(lambda x: len(x.split('.')))
    df_calc = df_calc.sort_values(by='Livello', ascending=False)
    
    for index, row in df_calc.iterrows():
        uid = str(row['ID_WBS'])
        if not row['Is_Leaf']:
            discendenti = df_calc[df_calc['ID_WBS'].astype(str).str.startswith(uid + '.') & df_calc['Is_Leaf']]
            if not discendenti.empty:
                df_calc.at[index, 'BAC_Budget'] = discendenti['BAC_Budget'].sum()
                df_calc.at[index, 'AC_Costo_Reale'] = discendenti['AC_Costo_Reale'].sum()
                
                inizio_min = pd.to_datetime(discendenti['Inizio_Previsto']).min()
                fine_max = pd.to_datetime(discendenti['Fine_Prevista']).max()
                if pd.notna(inizio_min): df_calc.at[index, 'Inizio_Previsto'] = inizio_min.date()
                if pd.notna(fine_max): df_calc.at[index, 'Fine_Prevista'] = fine_max.date()
                
                inizio_eff_min = pd.to_datetime(discendenti['Inizio_Effettivo']).min()
                fine_eff_max = pd.to_datetime(discendenti['Fine_Effettiva']).max()
                if pd.notna(inizio_eff_min): df_calc.at[index, 'Inizio_Effettivo'] = inizio_eff_min.date()
                if pd.notna(fine_eff_max): df_calc.at[index, 'Fine_Effettiva'] = fine_eff_max.date()
                
                tot_bac = discendenti['BAC_Budget'].sum()
                if tot_bac > 0:
                    df_calc.at[index, '%_Completamento'] = (discendenti['BAC_Budget'] * discendenti['%_Completamento']).sum() / tot_bac
                else:
                    df_calc.at[index, '%_Completamento'] = discendenti['%_Completamento'].mean()
                    
    df_calc['sort_key'] = df_calc['ID_WBS'].astype(str).apply(lambda x: '.'.join([p.zfill(5) for p in x.split('.')]))
    df_calc = df_calc.sort_values(by='sort_key').drop(columns=['sort_key', 'Is_Leaf', 'Livello']).reset_index(drop=True)
    return df_calc

def modifica_struttura(id_target, azione):
    df = st.session_state.wbs_data.copy()
    
    def get_sort_key(wbs_id):
        return [int(x) if x.isdigit() else x for x in str(wbs_id).split('.')]
    
    df['sort_key'] = df['ID_WBS'].apply(get_sort_key)
    df = df.sort_values(by='sort_key').reset_index(drop=True)
    df['Livello'] = df['ID_WBS'].apply(lambda x: len(str(x).split('.')))
    
    if azione == 'elimina':
        mask = (df['ID_WBS'].astype(str) == id_target) | (df['ID_WBS'].astype(str).str.startswith(f"{id_target}."))
        df = df[~mask].reset_index(drop=True) 
        
        if df.empty:
            df = pd.DataFrame([{'ID_WBS': '1', 'Attività': 'Progetto Principale', 'BAC_Budget': 0.0, '%_Completamento': 0.0, 'AC_Costo_Reale': 0.0, 'Livello': 1}])
            st.session_state.wbs_data = df.drop(columns=['Livello'])
            st.session_state['tracker_id'] = None
            for k in list(st.session_state.keys()):
                if k.startswith("editor_wbs_"):
                    del st.session_state[k]
            if azione != 'rinumera':
                st.rerun()
            
    elif azione in ['su', 'giu', 'destra', 'sinistra']:
        ids = df['ID_WBS'].astype(str).tolist()
        if id_target not in ids: return
        idx = ids.index(id_target)
        livello_target = df.at[idx, 'Livello']
        
        end_idx = idx + 1
        while end_idx < len(df) and df.at[end_idx, 'Livello'] > livello_target:
            end_idx += 1
        blocco_target = list(range(idx, end_idx))
        
        if azione == 'destra':
            if idx > 0 and df.at[idx - 1, 'Livello'] >= livello_target: df.loc[blocco_target, 'Livello'] += 1
        elif azione == 'sinistra':
            if livello_target > 1: df.loc[blocco_target, 'Livello'] -= 1
        elif azione == 'su':
            prev_idx = idx - 1
            while prev_idx >= 0 and df.at[prev_idx, 'Livello'] > livello_target: prev_idx -= 1
            if prev_idx >= 0 and df.at[prev_idx, 'Livello'] == livello_target:
                blocco_prev = list(range(prev_idx, idx))
                new_order = list(range(len(df)))
                new_order[prev_idx:end_idx] = blocco_target + blocco_prev
                df = df.iloc[new_order].reset_index(drop=True)
        elif azione == 'giu':
            next_idx = end_idx
            if next_idx < len(df) and df.at[next_idx, 'Livello'] == livello_target:
                next_end = next_idx + 1
                while next_end < len(df) and df.at[next_end, 'Livello'] > livello_target: next_end += 1
                blocco_next = list(range(next_idx, next_end))
                new_order = list(range(len(df)))
                new_order[idx:next_end] = blocco_next + blocco_target
                df = df.iloc[new_order].reset_index(drop=True)

    nuovi_id = []
    counters = {} 
    
    for idx, row in df.iterrows():
        liv = row['Livello']
        if idx == 0: liv = 1
        else:
            prev_liv = df.at[idx-1, 'Livello']
            if liv > prev_liv + 1: liv = prev_liv + 1 
                
        df.at[idx, 'Livello'] = liv
        counters[liv] = counters.get(liv, 0) + 1
        for k in list(counters.keys()):
            if k > liv: counters[k] = 0 
                
        nuovo_id = ".".join([str(counters[i]) for i in range(1, liv + 1)])
        nuovi_id.append(nuovo_id)
        
    old_ids = df['ID_WBS'].astype(str).tolist()
    mapping = dict(zip(old_ids, nuovi_id))
    
    def aggiorna_preds(val):
        if not val or pd.isna(val) or str(val).strip() in ['', 'None', 'nan']: return val
        preds = [p.strip() for p in str(val).split(',')]
        new_preds = []
        for p in preds:
            parts = p.split(' - ', 1)
            vecchio_id = parts[0].strip()
            nuovo_id = mapping.get(vecchio_id, vecchio_id)
            if len(parts) > 1: new_preds.append(f"{nuovo_id} - {parts[1]}")
            else: new_preds.append(nuovo_id)
        return ', '.join(new_preds)
        
    df['ID_WBS'] = nuovi_id
    if 'Predecessori' in df.columns:
        df['Predecessori'] = df['Predecessori'].apply(aggiorna_preds)
        
    df = df.drop(columns=['Livello', 'sort_key'])
    
    st.session_state.wbs_data = df.copy()
    st.session_state.wbs_data = aggiorna_gerarchia(st.session_state.wbs_data)
    
    if azione != 'elimina':
        st.session_state['tracker_id'] = mapping.get(id_target, id_target)
    else:
        st.session_state['tracker_id'] = None
        
    for k in list(st.session_state.keys()):
        if k.startswith("editor_wbs_"):
            del st.session_state[k]
            
    st.rerun()
    
def get_foglie(df):
    ids = df['ID_WBS'].astype(str).tolist()
    foglie = [uid for uid in ids if not any(other.startswith(uid + '.') for other in ids if other != uid)]
    return df[df['ID_WBS'].astype(str).isin(foglie)].copy()

def aggiorna_costi_reali():
    df_reg = st.session_state.registro_data.copy()
    if not df_reg.empty:
        df_reg['ID_WBS_calc'] = df_reg['Voce_WBS'].astype(str).apply(
            lambda x: str(x).split(' - ')[0].strip() if pd.notna(x) and str(x).strip() not in ['', 'None', 'nan'] else None
        )
        # Forza la conversione in numeri per evitare che stringhe o vuoti blocchino il calcolo
        df_reg['Importo_Netto'] = pd.to_numeric(df_reg['Importo_Netto'], errors='coerce').fillna(0.0)
        
        costi_raggruppati = df_reg.groupby('ID_WBS_calc')['Importo_Netto'].sum().reset_index()
        cost_map = dict(zip(costi_raggruppati['ID_WBS_calc'], costi_raggruppati['Importo_Netto']))
        wbs = st.session_state.wbs_data.copy()
        
        # Stacca gli spazi accidentali e inietta i costi nel database WBS
        wbs['AC_Costo_Reale'] = wbs['ID_WBS'].astype(str).str.strip().apply(lambda x: cost_map.get(x, 0.0))
        st.session_state.wbs_data = wbs

def calcola_evm(df, data_status):
    oggi = pd.to_datetime(data_status).date()
    df['BAC_Budget'] = pd.to_numeric(df['BAC_Budget'], errors='coerce').fillna(0.0)
    df['%_Completamento'] = pd.to_numeric(df['%_Completamento'], errors='coerce').fillna(0.0)
    df['AC_Costo_Reale'] = pd.to_numeric(df['AC_Costo_Reale'], errors='coerce').fillna(0.0)
    
    def calcola_pv(row):
        try:
            inizio_ts = pd.to_datetime(row['Inizio_Previsto'])
            fine_ts = pd.to_datetime(row['Fine_Prevista'])
            bac = float(row['BAC_Budget'])
            
            if pd.isna(inizio_ts) or pd.isna(fine_ts) or bac == 0: 
                return 0.0
            inizio = inizio_ts.date()
            fine = fine_ts.date()
            
            if oggi >= fine: return bac 
            if oggi <= inizio: return 0.0 
            
            giorni_totali = (fine - inizio).days
            giorni_trascorsi = (oggi - inizio).days
            if giorni_totali <= 0: return bac
            return bac * (giorni_trascorsi / giorni_totali)
        except Exception:
            return 0.0

    df['PV'] = df.apply(calcola_pv, axis=1)
    df['EV'] = df['BAC_Budget'] * (df['%_Completamento'] / 100.0)
    df['CV'] = df['EV'] - df['AC_Costo_Reale'] 
    df['SV'] = df['EV'] - df['PV']               
    
    df['SPI'] = df.apply(lambda x: (x['EV'] / x['PV']) if x['PV'] > 0 else (1.0 if x['EV']==0 else 1.1), axis=1)
    df['CPI'] = df.apply(lambda x: (x['EV'] / x['AC_Costo_Reale']) if x['AC_Costo_Reale'] > 0 else (1.0 if x['EV']==0 else 1.1), axis=1)

    df['EAC'] = df.apply(lambda x: x['BAC_Budget'] / x['CPI'] if x['CPI'] > 0 else x['BAC_Budget'], axis=1)
    df['ETC'] = df['EAC'] - df['AC_Costo_Reale']
    df['VAC'] = df['BAC_Budget'] - df['EAC']
    return df

def genera_dati_scurve(df_wbs, df_reg, data_status):
    oggi = pd.to_datetime(data_status).date()
    date_inizio = pd.to_datetime(df_wbs['Inizio_Previsto']).dropna().dt.date
    date_fine = pd.to_datetime(df_wbs['Fine_Prevista']).dropna().dt.date
    
    if date_inizio.empty or date_fine.empty:
        return None
        
    min_date = date_inizio.min()
    max_date = date_fine.max()
    date_range = pd.date_range(start=min_date, end=max_date)
    
    df_reg_calc = df_reg.copy()
    if not df_reg_calc.empty:
        df_reg_calc['Data'] = pd.to_datetime(df_reg_calc['Data'], errors='coerce').dt.date
        ac_daily = df_reg_calc.groupby('Data')['Importo_Netto'].sum().to_dict()
    else:
        ac_daily = {}
    
    dati = []
    cum_ac = 0.0
    
    for d_ts in date_range:
        d = d_ts.date()
        pv_giorno = 0.0
        ev_giorno = 0.0
        
        for _, row in df_wbs.iterrows():
            bac = float(row['BAC_Budget']) if pd.notna(row['BAC_Budget']) else 0.0
            ip = pd.to_datetime(row['Inizio_Previsto']).date() if pd.notna(row['Inizio_Previsto']) else None
            fp = pd.to_datetime(row['Fine_Prevista']).date() if pd.notna(row['Fine_Prevista']) else None
            if ip and fp and bac > 0:
                if d >= fp: 
                    pv_giorno += bac
                elif d > ip:
                    giorni_tot = (fp - ip).days
                    if giorni_tot > 0:
                        pv_giorno += bac * ((d - ip).days / giorni_tot)
                        
            if d <= oggi:
                ev_attuale = float(row['EV']) if 'EV' in row else 0.0
                ie = pd.to_datetime(row['Inizio_Effettivo']).date() if pd.notna(row['Inizio_Effettivo']) else ip
                if ie and ev_attuale > 0:
                    if d >= oggi:
                        ev_giorno += ev_attuale
                    elif d > ie:
                        giorni_lav = (oggi - ie).days
                        if giorni_lav > 0:
                            ev_giorno += ev_attuale * ((d - ie).days / giorni_lav)
        
        if d <= oggi:
            cum_ac += ac_daily.get(d, 0.0)
            ac_val = cum_ac
            ev_val = ev_giorno
        else:
            ac_val = None
            ev_val = None
            
        dati.append({'Data': d, 'PV (Valore Pianificato)': pv_giorno, 'EV (Valore Guadagnato)': ev_val, 'AC (Costo Reale)': ac_val})
    return pd.DataFrame(dati)

def calcola_cpm(df_wbs):
    df_wp = get_foglie(df_wbs)
    cpm_nodes = {}
    
    for _, row in df_wp.iterrows():
        node_id = str(row['ID_WBS']).strip()
        inizio = pd.to_datetime(row['Inizio_Previsto'], errors='coerce')
        fine = pd.to_datetime(row['Fine_Prevista'], errors='coerce')
        
        durata = max((fine - inizio).days + 1, 1) if pd.notna(inizio) and pd.notna(fine) else 1
        
        pred_val = str(row.get('Predecessori', '')).strip()
        preds = []
        if pred_val and pred_val.lower() not in ['none', 'nan', 'null']:
            for p in pred_val.split(','):
                p_id = p.split(' - ')[0].strip()
                if p_id.endswith('.0'): p_id = p_id[:-2]
                if p_id: preds.append(p_id)
        
        cpm_nodes[node_id] = {
            'durata': durata, 'preds': preds, 'succs': [],
            'ES': 0, 'EF': 0, 'LS': 0, 'LF': 0, 'slack': 0, 'is_critical': False
        }
        
    for n_id, data in cpm_nodes.items():
        for p_id in data['preds']:
            if p_id in cpm_nodes:
                cpm_nodes[p_id]['succs'].append(n_id)
                
    changed = True
    loop_counter = 0
    while changed and loop_counter < 1000:
        loop_counter += 1
        changed = False
        for n_id, data in cpm_nodes.items():
            max_ef = 0
            for p_id in data['preds']:
                if p_id in cpm_nodes:
                    max_ef = max(max_ef, cpm_nodes[p_id]['EF'])
            new_es = max_ef
            new_ef = new_es + data['durata']
            if new_es != data['ES'] or new_ef != data['EF']:
                data['ES'] = new_es; data['EF'] = new_ef; changed = True
                
    project_duration = max([data['EF'] for data in cpm_nodes.values()], default=0)
    for n_id, data in cpm_nodes.items():
        data['LF'] = project_duration
        data['LS'] = data['LF'] - data['durata']
        
    changed = True
    loop_counter = 0
    while changed and loop_counter < 1000:
        loop_counter += 1
        changed = False
        for n_id, data in cpm_nodes.items():
            min_ls = data['LF'] 
            if len(data['succs']) > 0:
                min_ls = min([cpm_nodes[s_id]['LS'] for s_id in data['succs'] if s_id in cpm_nodes])
            new_lf = min_ls
            new_ls = new_lf - data['durata']
            if new_lf != data['LF'] or new_ls != data['LS']:
                data['LF'] = new_lf; data['LS'] = new_ls; changed = True
                
    for n_id, data in cpm_nodes.items():
        data['slack'] = data['LS'] - data['ES']
        if data['slack'] <= 0:
            data['is_critical'] = True
            
    return cpm_nodes

# --- 3. ESECUZIONE CALCOLI INIZIALI ---
aggiorna_costi_reali()
st.session_state.wbs_data = aggiorna_gerarchia(st.session_state.wbs_data)
st.session_state.wbs_data = calcola_evm(st.session_state.wbs_data, pd.Timestamp.today().date())

# ==========================================
# IMPOSTAZIONE GRAFICA E LARGHEZZA DINAMICA
# ==========================================

# 1. Creiamo una "memoria" per la larghezza del pannello (default 1.2)
if 'pannello_sx' not in st.session_state:
    st.session_state.pannello_sx = 1.2

# 2. Le colonne ora "respirano" in base al valore del cursore
col_save, col_sviluppo = st.columns([st.session_state.pannello_sx, 10])

# ==========================================
# COLONNA DI SINISTRA (PANNELLO DI CONTROLLO)
# ==========================================
with col_save:
    # 3. Il cursore che comanda la larghezza in tempo reale!
    st.slider("↔️ Regola Pannello", min_value=1.0, max_value=5.0, step=0.1, key="pannello_sx", help="Trascina per allargare o restringere la colonna")

# ==========================================
# COLONNA DI SINISTRA (PANNELLO DI CONTROLLO)
# ==========================================
with col_save:
    st.caption("&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;PROGETTO")
    
    st.session_state.nome_progetto_attivo = st.text_area("Nome Progetto", value=st.session_state.nome_progetto_attivo, label_visibility="collapsed", height=100)
    
    st.markdown('<div class="btn-compatto">', unsafe_allow_html=True)
    
    # # --- 1. MEMORIA DI SESSIONE ---
    
    if st.button("💾 Salva", use_container_width=True):
        st.session_state.archivio_progetti[st.session_state.nome_progetto_attivo] = {
            "wbs": st.session_state.wbs_data.copy(),
            "obs": st.session_state.obs_data.copy(),
            "registro": st.session_state.registro_data.copy(),
            "capa": st.session_state.capa_data.copy(),
            "rischi": st.session_state.rischi_data.copy(),
            "sal": st.session_state.sal_data.copy(),
            "tickets": st.session_state.tickets_data.copy(),
            "memoria_burocratica": list(st.session_state.memoria_burocratica),
            "memoria_capa": list(st.session_state.memoria_capa),
            "memoria_ticket": list(st.session_state.memoria_ticket),
            
            # --- FIX: SCUDO ANTI-CRASH ---
            "conflitti_ignorati": list(st.session_state.get("conflitti_ignorati", [])),
            "imprevisti_ignorati": list(st.session_state.get("imprevisti_ignorati", [])),
            "varianti_ignorate": list(st.session_state.get("varianti_ignorate", [])),
            "rischi_rossi_ignorati": list(st.session_state.get("rischi_rossi_ignorati", []))
        }
        st.success("Salvato!")
        
    if st.button("📑 Duplica", use_container_width=True):
        nuovo_nome = f"{st.session_state.nome_progetto_attivo}_Copia"
        st.session_state.archivio_progetti[nuovo_nome] = {
            "wbs": st.session_state.wbs_data.copy(),
            "obs": st.session_state.obs_data.copy(),
            "registro": st.session_state.registro_data.copy(),
            "capa": st.session_state.capa_data.copy(),
            "rischi": st.session_state.rischi_data.copy(),
            "sal": st.session_state.sal_data.copy(),
            "tickets": st.session_state.tickets_data.copy(),
            "memoria_burocratica": list(st.session_state.memoria_burocratica),
            "memoria_capa": list(st.session_state.memoria_capa),
            "memoria_ticket": list(st.session_state.memoria_ticket),
            
            # --- FIX: SCUDO ANTI-CRASH ---
            "conflitti_ignorati": list(st.session_state.get("conflitti_ignorati", [])),
            "imprevisti_ignorati": list(st.session_state.get("imprevisti_ignorati", [])),
            "varianti_ignorate": list(st.session_state.get("varianti_ignorate", [])),
            "rischi_rossi_ignorati": list(st.session_state.get("rischi_rossi_ignorati", []))
        }
        st.session_state.nome_progetto_attivo = nuovo_nome
        st.rerun()

    if st.session_state.archivio_progetti:
        prog_selezionato = st.selectbox("Apri Progetto", options=list(st.session_state.archivio_progetti.keys()), label_visibility="collapsed")
        if st.button("📂 Apri Selezionato", use_container_width=True):
            st.session_state.wbs_data = st.session_state.archivio_progetti[prog_selezionato]["wbs"].copy()
            st.session_state.obs_data = st.session_state.archivio_progetti[prog_selezionato]["obs"].copy()
            st.session_state.registro_data = st.session_state.archivio_progetti[prog_selezionato]["registro"].copy()
            st.session_state.capa_data = st.session_state.archivio_progetti[prog_selezionato]["capa"].copy()
            st.session_state.rischi_data = st.session_state.archivio_progetti[prog_selezionato].get("rischi", pd.DataFrame()).copy()
            st.session_state.sal_data = st.session_state.archivio_progetti[prog_selezionato].get("sal", pd.DataFrame()).copy()
            st.session_state.conflitti_ignorati = st.session_state.archivio_progetti[prog_selezionato].get("conflitti_ignorati", []).copy()
            st.session_state.imprevisti_ignorati = st.session_state.archivio_progetti[prog_selezionato].get("imprevisti_ignorati", []).copy()
            st.session_state.varianti_ignorate = st.session_state.archivio_progetti[prog_selezionato].get("varianti_ignorate", []).copy()
            st.session_state.rischi_rossi_ignorati = st.session_state.archivio_progetti[prog_selezionato].get("rischi_rossi_ignorati", []).copy()
            st.session_state.tickets_data = st.session_state.archivio_progetti[prog_selezionato].get("tickets", pd.DataFrame()).copy()
            st.session_state.memoria_burocratica = set(st.session_state.archivio_progetti[prog_selezionato].get("memoria_burocratica", []))
            st.session_state.memoria_capa = set(st.session_state.archivio_progetti[prog_selezionato].get("memoria_capa", []))
            st.session_state.memoria_ticket = set(st.session_state.archivio_progetti[prog_selezionato].get("memoria_ticket", []))
            
            st.session_state.nome_progetto_attivo = prog_selezionato
            for k in list(st.session_state.keys()):
                if k.startswith("editor_wbs_"):
                    del st.session_state[k]
            st.rerun()

    if st.button("📄 Nuovo", use_container_width=True):
        st.session_state.nome_progetto_attivo = "Nuovo_Progetto"
        for key in ['wbs_data', 'obs_data', 'registro_data', 'capa_data', 'rischi_data', 'sal_data', 'tickets_data', 'memoria_burocratica', 'memoria_capa', 'memoria_ticket', 'conflitti_ignorati', 'imprevisti_ignorati', 'varianti_ignorate', 'rischi_rossi_ignorati']:
            if key in st.session_state:
                del st.session_state[key]
        for k in list(st.session_state.keys()):
            if k.startswith("editor_wbs_"):
                del st.session_state[k]
        st.rerun()
        
    # --- 2. ARCHIVIAZIONE SU PC (JSON) ---
    st.divider()
    st.caption("&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;ARCHIVIO PC")
    
    try:
        progetto_export = {
            "wbs": json.loads(st.session_state.wbs_data.to_json(orient="records", date_format="iso")),
            "obs": json.loads(st.session_state.obs_data.to_json(orient="records")),
            "registro": json.loads(st.session_state.registro_data.to_json(orient="records", date_format="iso")),
            "capa": json.loads(st.session_state.capa_data.to_json(orient="records", date_format="iso")),
            "rischi": json.loads(st.session_state.rischi_data.to_json(orient="records")),
            "sal": json.loads(st.session_state.sal_data.to_json(orient="records", date_format="iso")),
            "tickets": json.loads(st.session_state.tickets_data.to_json(orient="records", date_format="iso")),
            "memoria_burocratica": list(st.session_state.memoria_burocratica),
            "memoria_capa": list(st.session_state.memoria_capa),
            "memoria_ticket": list(st.session_state.memoria_ticket),
            
            # --- FIX: SCUDO ANTI-CRASH ---
            "conflitti_ignorati": list(st.session_state.get("conflitti_ignorati", [])),
            "imprevisti_ignorati": list(st.session_state.get("imprevisti_ignorati", [])),
            "varianti_ignorate": list(st.session_state.get("varianti_ignorate", [])),
            "rischi_rossi_ignorati": list(st.session_state.get("rischi_rossi_ignorati", []))
        }
        json_string = json.dumps(progetto_export, indent=4)
        
        st.download_button(
            label="⬇️ Scarica",
            data=json_string,
            file_name=f"{st.session_state.nome_progetto_attivo}.json",
            mime="application/json",
            use_container_width=True
        )
    except Exception as e:
        st.error(f"Errore: {e}")
    
    uploaded_file = st.file_uploader("📤 Carica da PC", type=['json'], label_visibility="collapsed")
    if uploaded_file is not None:
        if 'ultimo_file_caricato' not in st.session_state or st.session_state.ultimo_file_caricato != uploaded_file.file_id:
            try:
                dati_caricati = json.load(uploaded_file)
                
                df_wbs = pd.DataFrame(dati_caricati.get('wbs', []))
                if df_wbs.empty:
                    df_wbs = pd.DataFrame([{'ID_WBS': '1', 'Attività': 'Progetto Principale', 'BAC_Budget': 0.0, '%_Completamento': 0.0, 'AC_Costo_Reale': 0.0, 'ID_OBS_Assegnato': None, 'Predecessori': ''}])
                st.session_state.wbs_data = df_wbs
                
                df_obs = pd.DataFrame(dati_caricati.get('obs', []))
                if df_obs.empty:
                    df_obs = pd.DataFrame(columns=['ID_OBS', 'Ruolo', 'Risorsa', 'Tipo_Contratto', 'Note'])
                st.session_state.obs_data = df_obs
                
                df_reg = pd.DataFrame(dati_caricati.get('registro', []))
                if df_reg.empty:
                    df_reg = pd.DataFrame(columns=['Data', 'N_Doc', 'Fornitore', 'Voce_WBS', 'Importo_Netto', 'Descrizione'])
                st.session_state.registro_data = df_reg
                
                df_capa = pd.DataFrame(dati_caricati.get('capa', []))
                if df_capa.empty:
                    df_capa = pd.DataFrame(columns=['Data_Apertura', 'ID_WBS_Rif', 'Tipo_Azione', 'Descrizione', 'Responsabile_OBS', 'Stato', 'Rischio_Associato'])
                st.session_state.capa_data = df_capa

                df_rischi = pd.DataFrame(dati_caricati.get('rischi', []))
                if df_rischi.empty:
                    df_rischi = pd.DataFrame(columns=['ID_WBS_Rif', 'Descrizione_Rischio', 'Probabilità (1-5)', 'Impatto (1-5)', 'Stato'])
                st.session_state.rischi_data = df_rischi

                df_sal = pd.DataFrame(dati_caricati.get('sal', []))
                if df_sal.empty:
                    df_sal = pd.DataFrame(columns=['Data_Emissione', 'Descrizione_SAL', 'Importo_Euro', 'Stato_Pagamento'])
                st.session_state.sal_data = df_sal

                # --- INNESTO IMPORTAZIONE TICKETS ---
                df_tickets = pd.DataFrame(dati_caricati.get('tickets', []))
                if df_tickets.empty:
                    df_tickets = pd.DataFrame(columns=[
                    'ID_Ticket', 'ID_WBS_Rif', 'Autore', 'Data_Apertura',
                    'Tipologia', 'Descrizione', 'Stato', 'Risposta_RUP', 'Data_Chiusura',
                    'Variazione_Costi', 'Variazione_Tempi', 'Variante_Applicata'
                ])
                st.session_state.tickets_data = df_tickets
                # ------------------------------------

                st.session_state.conflitti_ignorati = dati_caricati.get('conflitti_ignorati', [])
                st.session_state.imprevisti_ignorati = dati_caricati.get('imprevisti_ignorati', [])
                st.session_state.varianti_ignorate = dati_caricati.get('varianti_ignorate', [])
                st.session_state.rischi_rossi_ignorati = dati_caricati.get('rischi_rossi_ignorati', [])
                st.session_state.memoria_burocratica = set(dati_caricati.get('memoria_burocratica', []))
                st.session_state.memoria_capa = set(dati_caricati.get('memoria_capa', []))
                st.session_state.memoria_ticket = set(dati_caricati.get('memoria_ticket', []))
                
                for col in ['Inizio_Previsto', 'Fine_Prevista', 'Inizio_Effettivo', 'Fine_Effettiva']:
                    if col in st.session_state.wbs_data.columns:
                        st.session_state.wbs_data[col] = pd.to_datetime(st.session_state.wbs_data[col], errors='coerce').apply(lambda x: x.date() if pd.notna(x) else None)
                if 'Data' in st.session_state.registro_data.columns:
                    st.session_state.registro_data['Data'] = pd.to_datetime(st.session_state.registro_data['Data'], errors='coerce').apply(lambda x: x.date() if pd.notna(x) else None)
                if 'Data_Apertura' in st.session_state.capa_data.columns:
                    st.session_state.capa_data['Data_Apertura'] = pd.to_datetime(st.session_state.capa_data['Data_Apertura'], errors='coerce').apply(lambda x: x.date() if pd.notna(x) else None)
                    
                st.session_state.wbs_data = aggiorna_gerarchia(st.session_state.wbs_data)
                st.session_state.nome_progetto_attivo = uploaded_file.name.replace(".json", "")
                st.session_state.ultimo_file_caricato = uploaded_file.file_id
                
                for k in list(st.session_state.keys()):
                    if k.startswith("editor_wbs_"):
                        del st.session_state[k]
                st.rerun() 
            except Exception as e:
                st.error(f"Errore critico durante la lettura: {e}")
                
    st.markdown('</div>', unsafe_allow_html=True)
    st.divider()
    st.caption("Versione 1.0")

   # ==========================================
    # 🤖 GIANFRY SUGGERISCE 🤖
    # ==========================================
    st.divider()
    st.markdown("#### 🤖 GIANFRY SUGGERISCE 🤖🤖🤖🤖")
    
    # Inizializziamo la memoria delle eccezioni consentite dal DL
    if 'conflitti_ignorati' not in st.session_state:
        st.session_state.conflitti_ignorati = []
    
    df_ai = get_foglie(st.session_state.wbs_data).copy()
    soglia_allerta = 0.95
    
    if df_ai.empty or df_ai['BAC_Budget'].sum() == 0:
        st.info("In attesa di dati.")
    else:
        # Rilevamento criticità EVM
        critici_costo = df_ai[df_ai['CPI'] < soglia_allerta]
        critici_tempo = df_ai[df_ai['SPI'] < soglia_allerta]
        
        # Rilevamento Rischi Attivi
        df_rischi_ai = st.session_state.rischi_data
        rischi_attivi = pd.DataFrame()
        if not df_rischi_ai.empty:
            rischi_attivi = df_rischi_ai[df_rischi_ai['Stato'].isin(['Attivo ▾', 'Monitorato ▾'])]

        # ========================================================
        # RADAR Gianfry: ALLARMI COMBINATI (RITARDO + RISCHI)
        # ========================================================
        df_wbs_radar = st.session_state.wbs_data.copy() if 'wbs_data' in st.session_state else pd.DataFrame()
        df_rischi_radar = st.session_state.rischi_data.copy() if 'rischi_data' in st.session_state else pd.DataFrame()

        allarmi_combinati = []
        oggi_radar = pd.Timestamp.today().normalize()
        
        if not df_wbs_radar.empty and not df_rischi_radar.empty:
            # 🎯 Identificazione flessibile (Bulletproof) delle colonne
            col_wbs_rischio = next((c for c in df_rischi_radar.columns if 'WBS' in c), None)
            col_stato_r = next((c for c in df_rischi_radar.columns if 'Stato' in c), None)
            col_impatto = next((c for c in df_rischi_radar.columns if 'Impatto' in c), None)
            col_desc = next((c for c in df_rischi_radar.columns if 'Descrizione' in c or 'Rischio' in c), None)
            
            if col_wbs_rischio and col_stato_r:
                for _, row_w in df_wbs_radar.iterrows():
                    wbs_id = str(row_w.get('ID_WBS', '')).strip()
                    
                    try:
                        completamento = float(row_w.get('%_Completamento', 0.0))
                    except:
                        completamento = 0.0
                        
                    fine_prev = pd.to_datetime(row_w.get('Fine_Prevista', None), errors='coerce', dayfirst=True)
                    
                    in_ritardo = False
                    # Se la data fine è superata e non siamo al 100%
                    if pd.notnull(fine_prev) and fine_prev < oggi_radar and completamento < 100.0:
                        in_ritardo = True
                        
                    if in_ritardo:
                        # Estrae "2.1.1" da "2.1.1 - demolizioni serramenti"
                        df_rischi_radar['wbs_clean'] = df_rischi_radar[col_wbs_rischio].astype(str).apply(lambda x: str(x).split(' - ')[0].strip())
                        rischi_associati = df_rischi_radar[df_rischi_radar['wbs_clean'] == wbs_id]
                        
                        for _, rischio in rischi_associati.iterrows():
                            stato_val = str(rischio.get(col_stato_r, '')).lower()
                            impatto_val = str(rischio.get(col_impatto, 'Non specificato'))
                            
                            # Se lo stato NON è chiuso (quindi è Attivo, Aperto, in lavorazione, ecc.)
                            if 'chiuso' not in stato_val and stato_val != 'nan' and stato_val != '':
                                desc_rischio = rischio.get(col_desc, 'Rischio generico') if col_desc else 'Rischio generico'
                                allarmi_combinati.append({
                                    'wbs': wbs_id,
                                    'nome_wbs': row_w.get('Attività', ''),
                                    'rischio': desc_rischio,
                                    'impatto': impatto_val
                                })

        # Stampa a schermo SOLO se ci sono criticità (Stealth mode)
        if allarmi_combinati:
            st.divider()
            st.markdown("### 🚨 Radar Gianfry: Criticità Combinate")
            for allarme in allarmi_combinati:
                st.error(f"⚠️ **INTERVENTO RUP:** L'attività **{allarme['wbs']} - {allarme['nome_wbs']}** è in ritardo cronico e ha un rischio attivo associato: *{allarme['rischio']}* (Impatto: {allarme['impatto']}). Intervenire o mitigare il rischio nel Tab 8.")
                
        # ===================================================
        # RILEVAMENTO CONFLITTI RISORSE (CON "IGNORE BUTTON")
        # ===================================================
        conflitti_risorse = []
        df_res = df_ai.dropna(subset=['Inizio_Previsto', 'Fine_Prevista']).copy()
        df_res = df_res[df_res['ID_OBS_Assegnato'].astype(str).str.strip().astype(bool)]
        df_res = df_res[~df_res['ID_OBS_Assegnato'].astype(str).isin(['None', 'nan'])]
        
        if not df_res.empty:
            df_res['Inizio_Previsto'] = pd.to_datetime(df_res['Inizio_Previsto'])
            df_res['Fine_Prevista'] = pd.to_datetime(df_res['Fine_Prevista'])
            
            for obs_val, group in df_res.groupby('ID_OBS_Assegnato'):
                tasks = group.to_dict('records')
                for i in range(len(tasks)):
                    for j in range(i + 1, len(tasks)):
                        t1, t2 = tasks[i], tasks[j]
                        inizio1, fine1 = t1['Inizio_Previsto'], t1['Fine_Prevista']
                        inizio2, fine2 = t2['Inizio_Previsto'], t2['Fine_Prevista']
                        
                        if (inizio1 <= fine2) and (inizio2 <= fine1):
                            # Creiamo una "Targa" univoca per questa sovrapposizione (es. "2.1_2.2")
                            conflitto_id = f"{t1['ID_WBS']}_{t2['ID_WBS']}"
                            
                            # Aggiungiamo all'allarme SOLO se non è stato ignorato dal DL
                            if conflitto_id not in st.session_state.conflitti_ignorati:
                                nome_risorsa = str(obs_val).split(' - ')[1] if ' - ' in str(obs_val) else str(obs_val)
                                sovrapposizione_inizio = max(inizio1, inizio2).strftime('%d/%m')
                                sovrapposizione_fine = min(fine1, fine2).strftime('%d/%m')
                                
                                conflitti_risorse.append({
                                    'ID_Univoco': conflitto_id,
                                    'Risorsa': nome_risorsa,
                                    'WBS1': f"{t1['ID_WBS']} ({t1['Attività']})",
                                    'WBS2': f"{t2['ID_WBS']} ({t2['Attività']})",
                                    'Date': f"dal {sovrapposizione_inizio} al {sovrapposizione_fine}"
                                })

        # --- GESTIONE DEGLI ALLARMI ---
        if critici_costo.empty and critici_tempo.empty and rischi_attivi.empty and not conflitti_risorse:
            st.success("✅ **Progetto in salute!** Nessuna azione richiesta.")
        else:
            # 1. Allarmi di Schedulazione (Tempi)
            if not critici_tempo.empty:
                for _, row in critici_tempo.iterrows():
                    with st.expander(f"⏳ WBS {row['ID_WBS']}: Ritardo"):
                        st.markdown(f"**Situazione:** Efficienza tempi al **{row['SPI']*100:.0f}%**.<br>**Azione:** Usa il Simulatore (*Tab 7*) per testare un *Crashing* oppure scala le date in *Tab 1*.", unsafe_allow_html=True)
            
            # 2. Allarmi Finanziari (Costi)
            if not critici_costo.empty:
                for _, row in critici_costo.iterrows():
                    with st.expander(f"💸 WBS {row['ID_WBS']}: Over-Budget"):
                        st.markdown(f"**Situazione:** Efficienza costi a **{row['CPI']:.2f}**.<br>**Azione:** Verifica le spese nel *Tab 6*. Se la spesa è irreversibile, apri una CAPA in *Tab 7*.", unsafe_allow_html=True)
                        
            # 3. Allarmi Rischio (Scudo)
            if not rischi_attivi.empty:
                with st.expander(f"⚠️ {len(rischi_attivi)} Rischi Attivi"):
                    st.markdown("**Situazione:** Capitale congelato nel Fondo Imprevisti.<br>**Azione:** Metti in atto le mitigazioni di cantiere e chiudi l'Azione in *Tab 7* per sbloccare i fondi.", unsafe_allow_html=True)
                    
            # 4. Allarmi Sovraccarico Risorse (Conflitti)
            if conflitti_risorse:
                with st.expander(f"👷 {len(conflitti_risorse)} Conflitti Risorse"):
                    st.markdown("**Situazione:** Una risorsa è stata assegnata a lavorazioni sovrapposte.")
                    for c in conflitti_risorse:
                        st.error(f"**{c['Risorsa']}** lavora su:\n* {c['WBS1']}\n* {c['WBS2']}\n*(Accavallamento: {c['Date']})*")
                        # IL BOTTONE PER IGNORARE!
                        if st.button("👁️ Consenti Sovrapposizione", key=f"ignora_{c['ID_Univoco']}", help="Nascondi questo allarme: accetto che lavorino in contemporanea."):
                            st.session_state.conflitti_ignorati.append(c['ID_Univoco'])
                            st.rerun()
                            
            # 5. NUOVO: Registro dei Conflitti Tollerati (Invisibili)
            if st.session_state.conflitti_ignorati:
                with st.expander(f"👁️ {len(st.session_state.conflitti_ignorati)} Conflitti Tollerati (Nascosti)"):
                    st.markdown("Hai autorizzato le seguenti sovrapposizioni:")
                    for conf_id in st.session_state.conflitti_ignorati:
                        wbs_split = conf_id.split('_')
                        st.markdown(f"- WBS **{wbs_split[0]}** + WBS **{wbs_split[1]}**")
                        if st.button("🔄 Ripristina Allarme", key=f"ripristina_{conf_id}"):
                            st.session_state.conflitti_ignorati.remove(conf_id)
                            st.rerun()
            
        # ===================================================
        # DA IMPREVISTI A RISCHI (SUGGERIMENTO)
        # ===================================================
        if 'imprevisti_ignorati' not in st.session_state:
            st.session_state.imprevisti_ignorati = []
            
        if isinstance(st.session_state.imprevisti_ignorati, set):
            st.session_state.imprevisti_ignorati = list(st.session_state.imprevisti_ignorati)

        if 'tickets_data' in st.session_state and not st.session_state.tickets_data.empty:
            if 'Tipologia' in st.session_state.tickets_data.columns and 'Stato' in st.session_state.tickets_data.columns:
                
                # Filtriamo i ticket "Imprevisto" che sono in attesa
                imprevisti = st.session_state.tickets_data[
                    (st.session_state.tickets_data['Tipologia'] == 'Segnalazione Imprevisto/Ostacolo') & 
                    (st.session_state.tickets_data['Stato'] == 'In attesa ⏳')
                ]
                
                for _, ticket in imprevisti.iterrows():
                    t_id = str(ticket.get('ID_Ticket', ''))
                    wbs_rif = str(ticket.get('ID_WBS_Rif', ''))
                    
                    if t_id.strip() != "":
                        if t_id in st.session_state.imprevisti_ignorati:
                            # FIX 4: Se Ignorato, lascia il riquadro ma lo colora di verde (Success)
                            with st.expander(f"✅ WBS {wbs_rif}: Imprevisto {t_id} (Risolto)"):
                                st.success(f"La segnalazione {t_id} è stata contrassegnata come 'Non è un rischio'. L'allarme è silenziato.")
                                if st.button("🔄 Ripristina Allarme", key=f"ripristina_imp_{t_id}"):
                                    st.session_state.imprevisti_ignorati.remove(t_id)
                                    st.rerun()
                        else:
                            # Altrimenti mostra l'allarme standard
                            with st.expander(f"🧩 WBS {wbs_rif}: Imprevisto Segnalato", expanded=True):
                                st.markdown(f"Dal campo è arrivata una **Segnalazione di Imprevisto/Ostacolo** (Ticket {t_id}).")
                                st.warning("💡 **Suggerimento:** Vai nel Tab 8 per mappare questo ostacolo come Rischio e calcolare l'impatto sul Fondo Imprevisti.")
                                
                                if st.button("👁️ Non è un rischio (Ignora)", key=f"ignora_imp_{t_id}"):
                                    st.session_state.imprevisti_ignorati.append(t_id)
                                    st.rerun()
        # ===================================================

        # ===================================================
        # DA VARIANTI APPROVATE A RISCHI (SUGGERIMENTO)
        # ===================================================
        if 'varianti_ignorate' not in st.session_state:
            st.session_state.varianti_ignorate = []
            
        if isinstance(st.session_state.varianti_ignorate, set):
            st.session_state.varianti_ignorate = list(st.session_state.varianti_ignorate)

        if 'tickets_data' in st.session_state and not st.session_state.tickets_data.empty:
            if 'Tipologia' in st.session_state.tickets_data.columns and 'Stato' in st.session_state.tickets_data.columns:
                
                # Filtriamo SOLO i ticket "Variante" che sono stati "Approvati"
                varianti_approvate = st.session_state.tickets_data[
                    (st.session_state.tickets_data['Tipologia'] == 'Richiesta di Variante') & 
                    (st.session_state.tickets_data['Stato'] == 'Approvato ✅')
                ]
                
                for _, ticket in varianti_approvate.iterrows():
                    t_id = str(ticket.get('ID_Ticket', ''))
                    wbs_rif = str(ticket.get('ID_WBS_Rif', ''))
                    
                    if t_id.strip() != "":
                        if t_id in st.session_state.varianti_ignorate:
                            # Se Ignorato, lascia il riquadro verde di conferma
                            with st.expander(f"✅ WBS {wbs_rif}: Variante {t_id} (Rischi Gestiti)"):
                                st.success(f"Hai confermato che la Variante {t_id} non comporta nuovi rischi o è già stata gestita nel Tab 8.")
                                if st.button("🔄 Ripristina Allarme", key=f"ripristina_var_{t_id}"):
                                    st.session_state.varianti_ignorate.remove(t_id)
                                    st.rerun()
                        else:
                            # Altrimenti mostra l'allarme standard
                            with st.expander(f"⚠️ WBS {wbs_rif}: Variante Approvata", expanded=True):
                                st.markdown(f"La **Variante {t_id}** è stata ufficialmente approvata dal RUP.")
                                st.warning("💡 **Gianfry suggerisce:** Questa variante comporta dei nuovi rischi per il cantiere? Ricordati di gestirli nel **Tab 8 - Matrice dei Rischi**.")
                                
                                if st.button("👁️ Rischi gestiti / Nessun rischio (Ignora)", key=f"ignora_var_{t_id}"):
                                    st.session_state.varianti_ignorate.append(t_id)
                                    st.rerun()
        # ===================================================

# ==========================================
# COLONNA DI DESTRA (IL MOTORE DELL'APP)
# ==========================================
with col_sviluppo:

    st.markdown("""
        <style>
            .block-container {
                padding-top: 1rem;
                padding-bottom: 0rem;
            }
        </style>
    """, unsafe_allow_html=True)

    col_logo, col_title = st.columns([1, 9], vertical_alignment="center")

    with col_logo:
        try:
            with open("logo_base64.txt", "r") as f:
                logo_str = f.read().strip()
            st.image(base64.b64decode(logo_str), width=120)
        except Exception:
            pass

    with col_title:
        st.title("Project Workflow & EVM Controller")

    # ==========================================
    # RETRO-COMPATIBILITÀ TICKETS E CAPA (Schema Update)
    # ==========================================
    if 'tickets_data' in st.session_state:
        if 'Variazione_Costi' not in st.session_state.tickets_data.columns:
            st.session_state.tickets_data['Variazione_Costi'] = None
        if 'Variazione_Tempi' not in st.session_state.tickets_data.columns:
            st.session_state.tickets_data['Variazione_Tempi'] = None
        if 'Variante_Applicata' not in st.session_state.tickets_data.columns:
            st.session_state.tickets_data['Variante_Applicata'] = False
            
    if 'capa_data' in st.session_state:
        if 'Costo_Intervento' not in st.session_state.capa_data.columns:
            st.session_state.capa_data['Costo_Intervento'] = 0.0
        if 'Giorni_Intervento' not in st.session_state.capa_data.columns:
            st.session_state.capa_data['Giorni_Intervento'] = 0
        if 'Costo_Scaricato' not in st.session_state.capa_data.columns:
            st.session_state.capa_data['Costo_Scaricato'] = False
            
    # ==========================================
    # 🦎 BANNER FISSO: GIANFRY ADVISOR (CAROSELLO)
    # ==========================================
    
    # 1. Inizializza la "memoria" per scorrere i messaggi
    if 'gianfry_idx' not in st.session_state:
        st.session_state.gianfry_idx = 0
    
    # 2. Il Motore che raccoglie TUTTI i consigli attivi (genera una lista)
    def ottieni_consigli_gianfry():
        consigli = []
        
        # QUI INSERIREMO LE VERE REGOLE COLLEGATE AI TUOI DATI
        # (Per ora metto 3 esempi fittizi per farti vedere come scorrono)
        consigli.append("👋 Ciao! Sono Gianfry. Ricordati di compilare tutte le date nel Tab 1 (WBS).")
        consigli.append("💰 Suggerimento: Hai attività con budget pari a zero. Verifica i costi per far girare l'EVM.")
        consigli.append("⏳ Attenzione: Il progetto è in leggero ritardo (SPI sotto lo 0.95). Controlla il Gantt.")
        
        # Se il sistema non trova nessun problema:
        if len(consigli) == 0:
            consigli.append("✅ Tutto perfetto! Nessuna criticità rilevata. Il cantiere procede a gonfie vele.")
            
        return consigli
    
    # 3. Interfaccia Visiva del Banner
    messaggi = ottieni_consigli_gianfry()
    tot_msg = len(messaggi)
    
    # Sicurezza: se correggi un errore e i messaggi diminuiscono, riavvolge il nastro
    if st.session_state.gianfry_idx >= tot_msg:
        st.session_state.gianfry_idx = 0
    
    # Impaginazione: Freccia SX | Messaggio Centrale | Freccia DX
    col_sx, col_centro, col_dx = st.columns([1, 10, 1], gap="small")
    
    with col_sx:
        # Freccia Sinistra (si disabilita se c'è un solo messaggio)
        if st.button("◀", key="g_prev", use_container_width=True, disabled=(tot_msg <= 1)):
            st.session_state.gianfry_idx = (st.session_state.gianfry_idx - 1) % tot_msg
    
    with col_centro:
        # Il Box colorato con il messaggio
        st.info(f"**Gianfry Consiglia ({st.session_state.gianfry_idx + 1}/{tot_msg}):** {messaggi[st.session_state.gianfry_idx]}")
    
    with col_dx:
        # Freccia Destra
        if st.button("▶", key="g_next", use_container_width=True, disabled=(tot_msg <= 1)):
            st.session_state.gianfry_idx = (st.session_state.gianfry_idx + 1) % tot_msg
    
    st.divider() # Una bella linea grigia di separazione prima di iniziare con i Tab
    
    # ==========================================
    
    # --- CREAZIONE TAB ---
    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8, tab9, tab10 = st.tabs([
        "🗂️ 1-WBS (Lavorazioni)", 
        "👥 2-OBS (Risorse)", 
        "🕸️ 3-Nodi & Matrice", 
        "📅 4-Cronoprogramma", 
        "📈 5-Earned Value & Cash Flow",
        "🧾 6-Reg. Contabile",
        "🛠️ 7-Direzione & CAPA",
        "⚠️ 8-Matrice Rischi",
        "📩 9-Varianti e Comunicazioni",
        "📚 10-Guida & Glossario"
    ])
        
    # --- TAB 1: SETUP WBS ---
    with tab1:
        st.header("WBS - Work Breakdown Structure")

        # ======================================================
        # ALLERTA FINANZIARIA GIANFRY: VARIANTI DA QUANTIFICARE
        if 'tickets_data' in st.session_state and not st.session_state.tickets_data.empty:
            df_t_alert = st.session_state.tickets_data
            if 'Variazione_Costi' in df_t_alert.columns and 'Variazione_Tempi' in df_t_alert.columns:
                
                # Conversione sicura per scovare le celle "" (vuote)
                costi_vuoti = pd.to_numeric(df_t_alert['Variazione_Costi'], errors='coerce').isna()
                tempi_vuoti = pd.to_numeric(df_t_alert['Variazione_Tempi'], errors='coerce').isna()
                
                # LOGICA CORRETTA: Suona solo se sono vuoti ENTRAMBI
                da_compilare = df_t_alert[
                    (df_t_alert['Stato'] == 'Approvato ✅') & 
                    (df_t_alert['Tipologia'] == 'Richiesta di Variante') & 
                    (costi_vuoti & tempi_vuoti)
                ]
                
                if not da_compilare.empty:
                    wbs_sospese = da_compilare['ID_WBS_Rif'].astype(str).tolist()
                    st.error(f"🚨 **ALLERTA FINANZIARIA (GIANFRY):** Hai approvato una Variante per le WBS **{', '.join(wbs_sospese)}** senza quantificarne l'impatto! Vai nel Tab 9 e compila 'Variazione Costi' o 'Variazione Tempi' per ricalibrare il motore EVM.")
        # ======================================================
        
        # ======================================================
        # INIZIO INNESTO: MEMORIA PERSISTENTE (AMMINISTRATIVA E QUALITÀ)
        if 'memoria_burocratica' not in st.session_state:
            st.session_state.memoria_burocratica = set()
        if 'memoria_capa' not in st.session_state:
            st.session_state.memoria_capa = set()
        if 'memoria_ticket' not in st.session_state:
            st.session_state.memoria_ticket = set()

        # Calcolo preventivo delle WBS bloccate per usarle nei controlli in tempo reale
        wbs_bloccate_capa = []
        if 'capa_data' in st.session_state and not st.session_state.capa_data.empty:
            capa_attive = st.session_state.capa_data[st.session_state.capa_data['Stato'].isin(['Aperto ▾', 'In Lavorazione ▾'])]
            if not capa_attive.empty:
                wbs_bloccate_capa = capa_attive['ID_WBS_Rif'].astype(str).apply(lambda x: x.split(' - ')[0].strip()).unique().tolist()
        
        wbs_bloccate_ticket = []
        if 'tickets_data' in st.session_state and not st.session_state.tickets_data.empty:
            ticket_attesi = st.session_state.tickets_data[st.session_state.tickets_data['Stato'] == 'In attesa ⏳']
            if not ticket_attesi.empty:
                wbs_bloccate_ticket = ticket_attesi['ID_WBS_Rif'].astype(str).str.strip().unique().tolist()

        # Pulizia automatica intelligente: se il ticket o la CAPA vengono chiusi dal RUP 
        # negli altri tab, l'avviso sparisce d'ufficio da qui senza dover fare nulla!
        st.session_state.memoria_capa = {w for w in st.session_state.memoria_capa if w in wbs_bloccate_capa}
        st.session_state.memoria_ticket = {w for w in st.session_state.memoria_ticket if w in wbs_bloccate_ticket}

        # DISPLAY DEGLI AVVISI IN FILA (Senza sovrapporsi, impaginati a colonna)
        if st.session_state.memoria_burocratica:
            st.error(f"🛑 **BLOCCO AMMINISTRATIVO:** Hai tentato di avviare le lavorazioni **{', '.join(st.session_state.memoria_burocratica)}** senza le dovute autorizzazioni. Ottieni il documento e metti la spunta su 'Vincolo Assolto'.")
            
        if st.session_state.memoria_ticket:
            st.warning(f"⏳ **CANCELLO SOSPESIVO:** Hai tentato di chiudere al 100% le WBS **{', '.join(st.session_state.memoria_ticket)}**. Sono state bloccate al 99% perché c'è un Ticket di Variante in attesa di approvazione RUP.")
            
        if st.session_state.memoria_capa:
            st.warning(f"🚧 **BLOCCO QUALITÀ:** Hai tentato di chiudere al 100% le WBS **{', '.join(st.session_state.memoria_capa)}**. Sono state bloccate al 99% perché c'è un'azione correttiva (CAPA) aperta nel Tab 7.")
        # FINE INNESTO MEMORIA PERSISTENTE
        # ======================================================
        
        st.markdown('*I numeri ID sono bloccati per garantire l\'integrità. Usa i pulsanti sotto ogni capitolo per spostare e rientrare le voci.*')
        
        st.subheader("🔀 Organizzatore Capitoli")
        c_sel, c_btn1, c_btn2, c_btn3 = st.columns([3, 1, 1, 1])
        
        df_padri = st.session_state.wbs_data[~st.session_state.wbs_data['ID_WBS'].astype(str).str.contains(r'\.')]
        lista_capitoli = list(df_padri['ID_WBS'].astype(str) + " - " + df_padri['Attività'].astype(str))
        
        # Inseguimento voce Capitolo
        tracker = st.session_state.get('tracker_id', None)
        idx_padre = 0
        if tracker:
            for i, opz in enumerate(lista_capitoli):
                if opz.split(' - ')[0] == tracker:
                    idx_padre = i
                    break
                    
        nodo_scelto = c_sel.selectbox("Seleziona Capitolo", options=lista_capitoli, index=idx_padre, label_visibility="collapsed")
        
        if nodo_scelto:
            id_scelto = nodo_scelto.split(' - ')[0]
            if c_btn1.button("⬆️ Su", use_container_width=True): modifica_struttura(id_scelto, 'su')
            if c_btn2.button("⬇️ Giù", use_container_width=True): modifica_struttura(id_scelto, 'giu')
            if c_btn3.button("🗑️ Elimina", use_container_width=True): modifica_struttura(id_scelto, 'elimina')
                
        st.divider()
        
        # ---  AGGIORNAMENTO RETROATTIVO COLONNE MANCANTI ---
        if 'Vincolo_Burocratico' not in st.session_state.wbs_data.columns:
            st.session_state.wbs_data['Vincolo_Burocratico'] = 'Nessuno'
    
        if 'Vincolo_Assolto' not in st.session_state.wbs_data.columns:
            st.session_state.wbs_data['Vincolo_Assolto'] = False
        # -----------------------------------------------------------
        
        df = st.session_state.wbs_data.copy()
        
        for col in ['Inizio_Previsto', 'Fine_Prevista', 'Inizio_Effettivo', 'Fine_Effettiva']:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors='coerce').dt.date
                
        df['Durata_Prevista (gg)'] = (pd.to_datetime(df['Fine_Prevista']) - pd.to_datetime(df['Inizio_Previsto'])).dt.days + 1
        
        radici = df[~df['ID_WBS'].astype(str).str.contains(r'\.')]
        df_aggiornato = pd.DataFrame()
        
        lista_obs_dropdown = [""] + [f"{row['ID_OBS']} - {row['Risorsa']}" for _, row in st.session_state.obs_data.iterrows() if pd.notna(row['ID_OBS'])]
        lista_wbs_dropdown = [""] + [f"{row['ID_WBS']} - {row['Attività']}" for _, row in df.iterrows() if pd.notna(row['ID_WBS'])]
        
        for idx_riga, radice in radici.iterrows():
            id_radice = str(radice['ID_WBS'])
            discendenti = df[df['ID_WBS'].astype(str).str.startswith(f"{id_radice}.")]
            tot_budget = radice['BAC_Budget']
            
            with st.expander(f"📁 {id_radice} - {radice['Attività']} (Budget Totale Raggruppato: € {tot_budget:,.2f})", expanded=True):
                
                colonne_bloccate = ["ID_WBS", "Durata_Prevista (gg)", "AC_Costo_Reale", "PV", "EV", "CV", "SV", "SPI", "CPI", "EAC", "ETC", "VAC"]
                colonne_bloccate = [col for col in colonne_bloccate if col in discendenti.columns]
                
                discendenti_modificati = st.data_editor(
                    discendenti,
                    key=f"editor_wbs_idx_{idx_riga}_id_{id_radice}",
                    num_rows="dynamic",
                    use_container_width=True,
                    hide_index=True,
                    disabled=colonne_bloccate, 
                    column_config={
                        "ID_WBS": st.column_config.TextColumn("ID WBS (Auto)"),
                        "Predecessori": st.column_config.SelectboxColumn("Predecessore ▾", options=lista_wbs_dropdown),
                        "ID_OBS_Assegnato": st.column_config.SelectboxColumn("Risorsa Assegnata ▾", options=lista_obs_dropdown),
                        "Inizio_Previsto": st.column_config.DateColumn("Inizio Previsto"),
                        "Fine_Prevista": st.column_config.DateColumn("Fine Prevista"),
                        "Inizio_Effettivo": st.column_config.DateColumn("Inizio Effettivo"),
                        "Fine_Effettiva": st.column_config.DateColumn("Fine Effettiva"),
                        "Vincolo_Burocratico": st.column_config.SelectboxColumn(
                            "🏛️ Vincolo Burocratico",
                            help="Seleziona l'autorizzazione necessaria per sbloccare l'inizio effettivo",
                            options=["Nessuno", "Delibera di giunta", "Deposito Genio Civile", "Autorizzazione Paesaggistica", "Validazione Progetto (RUP)", "Nomina CSE", "Nulla Osta Soprintendenza"],
                            default="Nessuno",
                            width="medium"
                        ),
                        "Vincolo_Assolto": st.column_config.CheckboxColumn(
                            "✅ Vincolo Assolto",
                            help="Spunta quando hai ottenuto il protocollo ufficiale",
                            default=False
                        ),
                    }
                )
                
                for i_row, row_mod in discendenti_modificati.iterrows():
                    val_id = str(row_mod['ID_WBS']).strip()
                    if val_id in ['', 'None', 'nan']:
                        discendenti_modificati.at[i_row, 'ID_WBS'] = f"{id_radice}.999{i_row}"
                
                # =====================================================================
                # --- INNESTO IN TEMPO REALE: CONTROLLI INCROCIATI ---
                allarmi_burocratici = []
                for i_row, row_mod in discendenti_modificati.iterrows():
                    nome_wbs = str(row_mod.get('ID_WBS', '')).strip()
                    
                    # 1. CANCELLO AMMINISTRATIVO (Blocco data Inizio)
                    vincolo = str(row_mod.get('Vincolo_Burocratico', '')).strip()
                    spunta = row_mod.get('Vincolo_Assolto', False)
                    sbloccato = True if str(spunta).strip().lower() in ['true', '1', 't', 'y', 'yes'] else False
                    
                    inizio = row_mod.get('Inizio_Effettivo')
                    has_inizio = False
                    if inizio is not None and pd.notna(inizio) and str(inizio).strip().lower() not in ['', 'nat', 'nan', 'none']:
                        has_inizio = True
                    
                    stringa_memoria = f"{nome_wbs} ({vincolo})"
                    
                    if vincolo not in ['Nessuno', '', 'nan', 'None']:
                        if not sbloccato and has_inizio:
                            discendenti_modificati.at[i_row, 'Inizio_Effettivo'] = pd.NaT 
                            allarmi_burocratici.append(nome_wbs)
                            st.session_state.memoria_burocratica.add(stringa_memoria)
                        elif sbloccato and stringa_memoria in st.session_state.memoria_burocratica:
                            st.session_state.memoria_burocratica.discard(stringa_memoria)
                            
                    # 2. BLOCCO QUALITÀ E TICKET (Blocco chiusura al 100%)
                    try:
                        completamento = float(row_mod.get('%_Completamento', 0))
                    except:
                        completamento = 0.0
                        
                    if completamento >= 100:
                        if nome_wbs in wbs_bloccate_ticket:
                            discendenti_modificati.at[i_row, '%_Completamento'] = 99.0
                            st.session_state.memoria_ticket.add(nome_wbs)
                            st.toast(f"⏳ CANCELLO SOSPESIVO: WBS {nome_wbs} bloccata al 99%.", icon='⏳')
                        elif nome_wbs in wbs_bloccate_capa:
                            discendenti_modificati.at[i_row, '%_Completamento'] = 99.0
                            st.session_state.memoria_capa.add(nome_wbs)
                            st.toast(f"🚧 BLOCCO QUALITÀ: WBS {nome_wbs} bloccata al 99%.", icon='🚧')
                    else:
                        # Se l'utente riporta manualmente la voce sotto il 100%, la memoria si ripulisce da sola
                        st.session_state.memoria_ticket.discard(nome_wbs)
                        st.session_state.memoria_capa.discard(nome_wbs)

                if allarmi_burocratici:
                    msg = f"Hai inserito l'Inizio per la WBS {', '.join(allarmi_burocratici)} senza spuntare il Vincolo Assolto. La data è stata annullata."
                    st.toast(msg, icon='🚨')
                # =====================================================================
 
                df_aggiornato = pd.concat([df_aggiornato, pd.DataFrame([radice]), discendenti_modificati], ignore_index=True)
                
                if not discendenti.empty:
                    st.markdown("↕️ **Sposta / Modifica Livello:**")
                    c_sel_int, c1, c2, c3, c4 = st.columns([3, 1.5, 1.5, 1, 1])
                    
                    opzioni_locali = list(discendenti['ID_WBS'].astype(str) + " - " + discendenti['Attività'].astype(str))
                    
                    # Inseguimento voce Sotto-Capitolo
                    idx_locale = 0
                    if tracker:
                        for i, opz in enumerate(opzioni_locali):
                            if opz.split(' - ')[0] == tracker:
                                idx_locale = i
                                break
                                
                    nodo_locale = c_sel_int.selectbox("Seleziona voce da muovere", options=opzioni_locali, index=idx_locale, key=f"sel_move_{id_radice}", label_visibility="collapsed")
                    
                    if nodo_locale:
                        id_loc = nodo_locale.split(' - ')[0]
                        if c1.button("⬅️ Rendi Padre", key=f"l_{id_radice}", use_container_width=True): modifica_struttura(id_loc, 'sinistra')
                        if c2.button("➡️ Rendi Figlio", key=f"r_{id_radice}", use_container_width=True): modifica_struttura(id_loc, 'destra')
                        if c3.button("⬆️ Su", key=f"u_{id_radice}", use_container_width=True): modifica_struttura(id_loc, 'su')
                        if c4.button("⬇️ Giù", key=f"d_{id_radice}", use_container_width=True): modifica_struttura(id_loc, 'giu')

        with st.form("aggiungi_padre"):
            st.write("Aggiungi un nuovo Capitolo Principale")
            c1, c2 = st.columns([4, 1])
            nuova_att = c1.text_input("Nome", placeholder="Es. Impianti Elettrici")
            if c2.form_submit_button("➕ Aggiungi Capitolo"):
                if nuova_att:
                    is_root_calc = ~st.session_state.wbs_data['ID_WBS'].astype(str).str.contains(r'\.')
                    nuovo_id = str(len(st.session_state.wbs_data[is_root_calc]) + 1)
                    nuova_riga = pd.DataFrame([{'ID_WBS': nuovo_id, 'Attività': nuova_att, 'BAC_Budget': 0.0, '%_Completamento': 0.0, 'AC_Costo_Reale': 0.0}])
                    st.session_state.wbs_data = pd.concat([st.session_state.wbs_data, nuova_riga], ignore_index=True)
                    modifica_struttura('1', 'rinumera') 

        st.divider()
        st.warning("⚠️ **Hai aggiunto nuove lavorazioni nelle tabelle?** Clicca il tasto qui sotto per far assegnare al sistema la numerazione definitiva e riallineare l'albero WBS.")
        if st.button("💾 SALVA INSERIMENTI E RICALCOLA ALBERO", type="primary", use_container_width=True, key="btn_salva_mega_wbs"):
            
            # --- 1. INIZIO INNESTO BLOCCHI QUALITÀ (CAPA + TICKETS) ---
            # A) Raccolta WBS bloccate da CAPA (Non-Conformità)
            wbs_bloccate_capa = []
            if not st.session_state.capa_data.empty and 'ID_WBS_Rif' in st.session_state.capa_data.columns:
                capa_attive = st.session_state.capa_data[st.session_state.capa_data['Stato'].isin(['Aperto ▾', 'In Lavorazione ▾'])]
                if not capa_attive.empty:
                    wbs_bloccate_capa = capa_attive['ID_WBS_Rif'].astype(str).apply(lambda x: x.split(' - ')[0].strip()).unique().tolist()
            
            # B) Raccolta WBS bloccate da TICKET (Cancello Sospensivo)
            wbs_bloccate_ticket = []
            if 'tickets_data' in st.session_state and not st.session_state.tickets_data.empty:
                ticket_attesi = st.session_state.tickets_data[st.session_state.tickets_data['Stato'] == 'In attesa ⏳']
                if not ticket_attesi.empty:
                    wbs_bloccate_ticket = ticket_attesi['ID_WBS_Rif'].astype(str).str.strip().unique().tolist()

            allarmi_capa = []
            allarmi_ticket = []
            
            for idx, row in df_aggiornato.iterrows():
                wbs_id = str(row.get('ID_WBS', '')).strip()
                try:
                    completamento = float(row.get('%_Completamento', 0))
                except:
                    completamento = 0.0
                    
                if completamento >= 100:
                    # La priorità va al Cancello Sospensivo (Ticket istituzionali)
                    if wbs_id in wbs_bloccate_ticket:
                        df_aggiornato.at[idx, '%_Completamento'] = 99.0
                        allarmi_ticket.append(wbs_id)
                    # Poi controlla le anomalie esecutive (CAPA di cantiere)
                    elif wbs_id in wbs_bloccate_capa:
                        df_aggiornato.at[idx, '%_Completamento'] = 99.0
                        allarmi_capa.append(wbs_id)
            # --- FINE INNESTO BLOCCHI QUALITÀ ---
            
            # --- INIZIO INNESTO: CANCELLO AMMINISTRATIVO ---
            allarmi_burocratici = []
            for idx, row in df_aggiornato.iterrows():
                vincolo = str(row.get('Vincolo_Burocratico', 'Nessuno')).strip()
                
                # 1. CATTURA IL BOOLEANO A PROVA DI BOMBA
                val_spunta = row.get('Vincolo_Assolto', False)
                if isinstance(val_spunta, bool):
                    sbloccato = val_spunta
                else:
                    # Se è un testo, lo converte. Se è vuoto o nan, diventa False.
                    sbloccato = str(val_spunta).strip().lower() in ['true', '1', 't', 'y', 'yes']
                
                # 2. CATTURA LA DATA IN MODO SICURO
                inizio_effettivo = row.get('Inizio_Effettivo', None)
                has_inizio = pd.notna(inizio_effettivo) and str(inizio_effettivo).strip().lower() not in ['none', 'nat', 'nan', '']
                
                # 3. IL BLOCCO
                if vincolo not in ['Nessuno', 'nan', '', 'None'] and not sbloccato and has_inizio:
                    # Usa pd.NaT invece di None per svuotare correttamente una colonna DateTime in Pandas
                    df_aggiornato.at[idx, 'Inizio_Effettivo'] = pd.NaT 
                    allarmi_burocratici.append(f"{row.get('ID_WBS')} ({vincolo})")
            
            if allarmi_burocratici:
                st.error(f"🛑 BLOCCO AMMINISTRATIVO: La data di 'Inizio Effettivo' delle seguenti lavorazioni è stata annullata perché il vincolo burocratico non è ancora stato assolto: {', '.join(allarmi_burocratici)}")
            # --- FINE INNESTO CANCELLO AMMINISTRATIVO ---
            
            # 2. Salvataggio del dataframe corretto
            st.session_state.wbs_data = df_aggiornato
            
            # 3. Pulizia della cache di Streamlit
            for k in list(st.session_state.keys()):
                if k.startswith("editor_wbs_"):
                    del st.session_state[k]
                    
            # 4. Ricalcolo struttura
            modifica_struttura('1', 'rinumera')
            
            # 5. Feedback a schermo
            bloccato_qualcosa = False
            
            if allarmi_capa:
                st.error(f"🚧 BLOCCO QUALITÀ: WBS {', '.join(allarmi_capa)} bloccate al 99% per CAPA aperte nel Tab 7.")
                bloccato_qualcosa = True
                
            if allarmi_ticket:
                st.error(f"⏳ CANCELLO SOSPESIVO: WBS {', '.join(allarmi_ticket)} bloccate al 99%. Attendi l'approvazione della variante nel Tab 9.")
                bloccato_qualcosa = True
                
            if not bloccato_qualcosa:
                st.success("✅ Dati salvati e albero ricalcolato!")
                import time
                time.sleep(1.0)
            else:
                import time
                time.sleep(3.5)
                
            st.rerun()

    # --- TAB 2: SETUP OBS ---
    with tab2:
        st.header("OBS - Organization Breakdown Structure")
        
        with st.expander("⚙️ Gestione Colonne Aggiuntive", expanded=False):
            c1, c2 = st.columns([3, 1])
            nuova_col = c1.text_input("Nome nuova colonna (es. Telefono, Qualifica, Partita IVA)")
            if c2.button("➕ Crea") and nuova_col:
                if nuova_col not in st.session_state.obs_data.columns:
                    st.session_state.obs_data[nuova_col] = "" 
                    st.rerun()
            
            colonne_custom = [c for c in st.session_state.obs_data.columns if c not in ['ID_OBS', 'Ruolo', 'Risorsa', 'Tipo_Contratto', 'Note']]
            if colonne_custom:
                st.divider()
                c3, c4, c5 = st.columns([2, 2, 1])
                col_da_modificare = c3.selectbox("Colonna da rinominare", options=colonne_custom)
                nuovo_nome = c4.text_input("Nuovo nome intestazione")
                if c5.button("✏️ Modifica") and nuovo_nome:
                    st.session_state.obs_data.rename(columns={col_da_modificare: nuovo_nome}, inplace=True)
                    st.rerun()

        edited_obs = st.data_editor(
            st.session_state.obs_data, 
            column_config={
                "Tipo_Contratto": st.column_config.SelectboxColumn("Tipo Contratto", options=["Appalto ▾", "Sub appalto ▾"])
            },
            num_rows="dynamic", use_container_width=True, hide_index=True
        )
        
        st.divider()
        if st.button("💾 CONFERMA E SALVA ANAGRAFICA (OBS)", type="primary", use_container_width=True):
            st.session_state.obs_data = edited_obs
            st.success("✅ Dati anagrafici salvati con successo!")
            st.rerun()
    
    # --- TAB 3: MATRICE E GRAFO A NODI ---
    with tab3:
        st.header("Percorso Logico - Work Packages e Percorso Critico")
        
        try:
            cpm_data = calcola_cpm(st.session_state.wbs_data)
            mostra_relazioni = st.toggle("👁️ Mostra Percorso Critico (CPM)", value=True)
            
            graph = graphviz.Digraph(engine='dot')
            graph.attr(rankdir='LR', ranksep='1.5', nodesep='0.8', splines='spline')
            graph.attr('node', fontname='Helvetica', fontsize='10', margin='0.2')
            
            # Scudo totale contro caratteri che fanno impazzire il motore grafico
            def pulisci_testo(testo):
                if pd.isna(testo) or testo is None: return ""
                t = str(testo)
                t = t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')
                t = t.replace('\n', '<BR/>') # Questo impedisce i crash se premi "invio" nelle descrizioni
                return t
                
            for _, row in st.session_state.obs_data.iterrows():
                ruolo = pulisci_testo(row.get('Ruolo', ''))
                risorsa = pulisci_testo(row.get('Risorsa', ''))
                
                if not ruolo and not risorsa: continue
                
                label_html = f"<<TABLE BORDER='0' CELLBORDER='0' CELLSPACING='2'>"
                label_html += f"<TR><TD><B>{ruolo}</B></TD></TR>"
                label_html += f"<TR><TD>({risorsa})</TD></TR>"
                
                colonne_custom = [col for col in st.session_state.obs_data.columns if col not in ['ID_OBS', 'Ruolo', 'Risorsa', 'Tipo_Contratto', 'Note']]
                for col in colonne_custom:
                    valore = pulisci_testo(row.get(col, ''))
                    if pd.notna(row.get(col)) and valore.strip() != "":
                        col_safe = pulisci_testo(col)
                        label_html += f"<TR><TD><FONT POINT-SIZE='9' COLOR='gray30'>{col_safe}: {valore}</FONT></TD></TR>"
                label_html += "</TABLE>>"
                
                obs_id = str(row.get('ID_OBS', '')).strip()
                if obs_id.endswith('.0'): obs_id = obs_id[:-2]
                
                if obs_id:
                    graph.node(f"OBS_{obs_id}", label=label_html, shape='rect', style='rounded,filled', fillcolor='#E1F5FE', color='#0288D1', penwidth='1.5')
                
            df_wp_reali = get_foglie(st.session_state.wbs_data)
            valid_wbs_ids = set(df_wp_reali['ID_WBS'].astype(str))
            
            if df_wp_reali.empty or len(valid_wbs_ids) == 0:
                graph.node("Vuoto", label="Nessuna lavorazione inserita.", shape="rect")
                
            # --- AGGIORNAMENTO GRAFO: INTEGRAZIONE RISCHI ---
            df_rischi = st.session_state.rischi_data
            
            for _, row in df_wp_reali.iterrows():
                wbs_id = str(row.get('ID_WBS', '')).strip()
                if not wbs_id or wbs_id in ['nan', 'None']: continue
                
                attivita = pulisci_testo(row.get('Attività', ''))
                
                # 1. Base Styling del CPM (Percorso Critico Temporale)
                wp_cpm = cpm_data.get(wbs_id, {})
                is_critical = wp_cpm.get('is_critical', False)
                margine = wp_cpm.get('slack', 0)
                
                if is_critical:
                    testo_margine = f"<FONT COLOR='#D32F2F'><B>Margine: {margine} gg</B></FONT>"
                    bordo_colore, spessore_bordo = '#D32F2F', '3.0'
                else:
                    testo_margine = f"<FONT COLOR='#388E3C'>Margine: {margine} gg</FONT>"
                    bordo_colore, spessore_bordo = '#388E3C', '1.5'

                # 2. Sovrascrittura Rischio (SOLO se Attivo o Monitorato)
                rischi_wbs = df_rischi[df_rischi['ID_WBS_Rif'].astype(str).str.startswith(wbs_id + " -")]
                rischi_attivi = rischi_wbs[rischi_wbs['Stato'].isin(['Attivo ▾', 'Monitorato ▾'])]
                
                punteggio_max = 0
                if not rischi_attivi.empty:
                    prob = pd.to_numeric(rischi_attivi['Probabilità (1-5)'], errors='coerce').fillna(0)
                    imp = pd.to_numeric(rischi_attivi['Impatto (1-5)'], errors='coerce').fillna(0)
                    punteggio_max = (prob * imp).max()

                # 3. Applicazione Alert Visivo
                tag_rischio = ""
                if punteggio_max >= 15:
                    bordo_colore, spessore_bordo = '#FF0000', '4.0' # Rosso Fuoco per emergenze
                    tag_rischio = " <FONT COLOR='#FF0000'><b>[💣 ALTO RISCHIO]</b></FONT>"
                elif punteggio_max >= 8:
                    if not is_critical: 
                        bordo_colore, spessore_bordo = '#FF9800', '2.5' # Arancione
                    tag_rischio = " <FONT COLOR='#FF9800'><b>[⚠️ RISCHIO]</b></FONT>"

                # 4. Parametri classici (Budget, Costi, ecc.)
                try: budget = float(str(row.get('BAC_Budget', 0)).replace(',', '.'))
                except: budget = 0.0
                try: costo_reale = float(str(row.get('AC_Costo_Reale', 0)).replace(',', '.'))
                except: costo_reale = 0.0
                try: completamento = float(str(row.get('%_Completamento', 0)).replace(',', '.'))
                except: completamento = 0.0
                
                inizio_val = pd.to_datetime(row.get('Inizio_Previsto'), errors='coerce')
                inizio_str = inizio_val.strftime('%d/%m/%Y') if pd.notna(inizio_val) else "N/D"
                fine_val = pd.to_datetime(row.get('Fine_Prevista'), errors='coerce')
                fine_str = fine_val.strftime('%d/%m/%Y') if pd.notna(fine_val) else "N/D"
                
                # Assemblaggio Nodo HTML
                wp_html = f"<<TABLE BORDER='0' CELLBORDER='0' CELLSPACING='4'>"
                wp_html += f"<TR><TD COLSPAN='2'><B>{wbs_id} - {attivita}</B>{tag_rischio}</TD></TR>"
                wp_html += f"<TR><TD ALIGN='LEFT'>Inizio: {inizio_str}</TD><TD ALIGN='RIGHT'>Fine: {fine_str}</TD></TR>"
                wp_html += f"<TR><TD ALIGN='LEFT'>Budget: {budget:,.0f} Euro</TD><TD ALIGN='RIGHT'>AC: {costo_reale:,.0f} Euro</TD></TR>"
                wp_html += f"<TR><TD ALIGN='LEFT'>Avanzamento: {completamento:.0f}%</TD><TD ALIGN='RIGHT'>{testo_margine}</TD></TR>"
                wp_html += "</TABLE>>"
                
                # Colore Barra Progressione
                if completamento >= 100:
                    stile, colore_sfondo = 'rounded,filled', '#C8E6C9' 
                elif completamento <= 0:
                    stile, colore_sfondo = 'rounded,filled', 'white'   
                else:
                    stile = 'rounded,striped'
                    quota_verde = max(0.01, min(0.99, completamento / 100.0))
                    colore_sfondo = f"#C8E6C9;{str(round(quota_verde, 3)).replace(',', '.')}:white"
                    
                graph.node(f"WBS_{wbs_id}", label=wp_html, shape='rect', style=stile, fillcolor=colore_sfondo, color=bordo_colore, penwidth=spessore_bordo)
                     
                obs_val = str(row.get('ID_OBS_Assegnato', '')).strip()
                if obs_val and obs_val.lower() not in ['none', 'nan', 'null']:
                    for o in [o.strip() for o in obs_val.split(',')]:
                        o_id = o.split(' - ')[0].strip()
                        if o_id.endswith('.0'): o_id = o_id[:-2]
                        if o_id: graph.edge(f"OBS_{o_id}", f"WBS_{wbs_id}", color='#757575', penwidth='1.5', arrowsize='0.8')
                            
                pred_val = str(row.get('Predecessori', '')).strip()
                if mostra_relazioni and pred_val and pred_val.lower() not in ['none', 'nan', 'null']:
                    for p in [p.strip() for p in pred_val.split(',')]:
                        p_id = p.split(' - ')[0].strip()
                        if p_id.endswith('.0'): p_id = p_id[:-2]
                        if p_id in valid_wbs_ids:
                            pred_is_critical = cpm_data.get(p_id, {}).get('is_critical', False)
                            col_cavo, stile_cavo, spes_cavo, freccia = ('#D32F2F', 'solid', '2.5', '1.0') if (is_critical and pred_is_critical) else ('#FF9800', 'dashed', '1.0', '0.6')
                            graph.edge(f"WBS_{p_id}", f"WBS_{wbs_id}", color=col_cavo, style=stile_cavo, penwidth=spes_cavo, arrowsize=freccia)

            # RENDERIZZAZIONE SICURA
            raw_svg = graph.pipe(format='svg').decode('utf-8')
            if '<svg' in raw_svg:
                svg_data = raw_svg[raw_svg.find('<svg'):]
                html_code = f"""
                <!DOCTYPE html><html><head>
                <script src="https://cdn.jsdelivr.net/npm/svg-pan-zoom@3.6.1/dist/svg-pan-zoom.min.js"></script>
                <style>body {{ margin: 0; overflow: hidden; background-color: #fafafa; }} #svg-container {{ width: 100vw; height: 100vh; display: flex; justify-content: center; align-items: center; }} svg {{ width: 100% !important; height: 100% !important; }}</style>
                </head><body><div id="svg-container">{svg_data}</div>
                <script>window.onload = function() {{ 
                    var s = document.querySelector('svg'); 
                    if (s) {{ 
                        s.setAttribute('id', 'grafo'); 
                        s.removeAttribute('width'); 
                        s.removeAttribute('height'); 
                        var panZoom = svgPanZoom('#grafo', {{ 
                            zoomEnabled: true, 
                            controlIconsEnabled: true, 
                            fit: false, 
                            center: true,
                            minZoom: 0.1,
                            maxZoom: 10,
                            mouseWheelZoomEnabled: true
                        }}); 
                    }} 
                }};</script>
                </body></html>
                """
                components.html(html_code, height=900)
            else:
                st.warning("Grafo generato ma vuoto.")
                
        except Exception as e:
            st.error(f"⚠️ ERRORE TECNICO DURANTE IL DISEGNO DEL GRAFO: {e}")
            
        st.divider()
        st.subheader("📖 Legenda del Grafo")
        col_leg1, col_leg2 = st.columns(2)
        with col_leg1:
            st.markdown("**NODI E FIGURE**\n* 🟦 **Riquadro Azzurro:** Risorsa/Ruolo (OBS)\n* 🟩 **Riquadro Verde:** Work Package (WBS)\n* 🟥 **Bordo Rosso Spesso:** Percorso Critico (Margine = 0 gg)")
        with col_leg2:
            st.markdown("**CAVI E COLLEGAMENTI**\n* 🔗 **Freccia Grigia Continua:** Assegnazione Risorsa\n* 🔀 **Freccia Arancione Tratteggiata:** Relazione logica\n* 🚨 **Freccia Rossa Spessa:** Flusso del Percorso Critico")

   # --- TAB 4: CRONOPROGRAMMA (GANTT) ---
    with tab4:
        st.header("Cronoprogramma & Scadenzario Amministrativo")
        
        # ========================================================
        # 1. GANTT CHART (IL CODICE ESISTENTE)
        # ========================================================
        st.subheader("📊 Diagramma di Gantt (EVM-Aware)")
        st.markdown("Le barre della fase esecutiva cambiano automaticamente colore in base alle performance reali (Indice SPI).")
        
        c1, c2, c3 = st.columns([1, 1, 1])
        vista = c1.selectbox("Seleziona Vista", ["Progetto (Baseline)", "Esecuzione (Esecutivo)", "Comparativa"])
        mostra_frecce = c2.toggle("🔗 Mostra Frecce Dipendenze", value=True)
        data_status_gantt = c3.date_input("📅 Data Rilevamento", value=pd.Timestamp.today().date())
        
        df_gantt = get_foglie(st.session_state.wbs_data).copy().dropna(subset=['Inizio_Previsto', 'Fine_Prevista'])
        
        if not df_gantt.empty:
            df_gantt['Inizio_Previsto'] = pd.to_datetime(df_gantt['Inizio_Previsto'])
            df_gantt['Fine_Prevista'] = pd.to_datetime(df_gantt['Fine_Prevista'])
            df_gantt['Inizio_Effettivo'] = pd.to_datetime(df_gantt['Inizio_Effettivo'])
            df_gantt['Fine_Effettiva'] = pd.to_datetime(df_gantt['Fine_Effettiva']).fillna(pd.to_datetime(data_status_gantt))
            
            cpm_data = calcola_cpm(st.session_state.wbs_data)
            df_evm_gantt = calcola_evm(get_foglie(st.session_state.wbs_data), data_status_gantt)
            
            fig = go.Figure()
            
            if vista in ["Progetto (Baseline)", "Comparativa"]:
                fig.add_trace(go.Bar(
                    x=(df_gantt['Fine_Prevista'] - df_gantt['Inizio_Previsto'] + pd.Timedelta(days=1)).dt.total_seconds() * 1000, 
                    y=df_gantt['ID_WBS'].astype(str) + " - " + df_gantt['Attività'], 
                    base=df_gantt['Inizio_Previsto'], 
                    orientation='h', name='Baseline (Pianificato)', width=0.4, 
                    marker=dict(color='rgba(30, 136, 229, 0.4)' if vista == "Comparativa" else '#1E88E5')
                ))
                
            if vista in ["Esecuzione (Esecutivo)", "Comparativa"]:
                df_esec = df_gantt.dropna(subset=['Inizio_Effettivo']).copy()
                if not df_esec.empty:
                    df_esec['ID_WBS'] = df_esec['ID_WBS'].astype(str).str.strip()
                    df_evm_clean = df_evm_gantt[['ID_WBS', 'SPI', '%_Completamento']].copy()
                    df_evm_clean['ID_WBS'] = df_evm_clean['ID_WBS'].astype(str).str.strip()
                    df_esec = df_esec.drop(columns=['SPI', '%_Completamento'], errors='ignore')
                    df_esec = df_esec.merge(df_evm_clean, on='ID_WBS', how='left')
                    
                    def colora_gantt(row):
                        spi = row['SPI']
                        if pd.isna(spi) or row['%_Completamento'] == 0: return '#9E9E9E'
                        if spi >= 1.0: return '#4CAF50'
                        if spi >= 0.90: return '#FF9800'
                        return '#D32F2F'
                        
                    colori_barre = df_esec.apply(colora_gantt, axis=1).tolist()
                    testi_hover = df_esec['SPI'].apply(lambda x: f"SPI: {x:.2f}" if pd.notna(x) else "").tolist()

                    fig.add_trace(go.Bar(
                        x=(df_esec['Fine_Effettiva'] - df_esec['Inizio_Effettivo'] + pd.Timedelta(days=1)).dt.total_seconds() * 1000, 
                        y=df_esec['ID_WBS'].astype(str) + " - " + df_esec['Attività'], 
                        base=df_esec['Inizio_Effettivo'], orientation='h', name='Esecutivo',
                        text=testi_hover, textposition='inside', insidetextanchor='middle',
                        textfont=dict(color='white', size=11, family='Arial Black'),
                        width=0.4 if vista == "Esecuzione (Esecutivo)" else 0.2, 
                        marker=dict(color=colori_barre)
                    ))
                    
                    fig.add_trace(go.Scatter(x=[None], y=[None], mode='markers', marker=dict(size=12, color='#4CAF50', symbol='square'), name='🟢 Puntuale/Anticipo (SPI ≥ 1)'))
                    fig.add_trace(go.Scatter(x=[None], y=[None], mode='markers', marker=dict(size=12, color='#FF9800', symbol='square'), name='🟠 Lieve Ritardo (SPI 0.9-1.0)'))
                    fig.add_trace(go.Scatter(x=[None], y=[None], mode='markers', marker=dict(size=12, color='#D32F2F', symbol='square'), name='🔴 Ritardo Grave (SPI < 0.9)'))
                    
            if mostra_frecce:
                for _, row in df_gantt.iterrows():
                    wbs_id = str(row['ID_WBS']).strip()
                    succ_y = wbs_id + " - " + str(row['Attività'])
                    if vista == "Esecuzione (Esecutivo)" and pd.notna(row['Inizio_Effettivo']): succ_start = row['Inizio_Effettivo']
                    else: succ_start = row['Inizio_Previsto']
                    
                    preds = str(row.get('Predecessori', '')).strip()
                    if preds and preds.lower() not in ['none', 'nan', 'null']:
                        for p in preds.split(','):
                            p_id = p.split(' - ')[0].strip()
                            if p_id.endswith('.0'): p_id = p_id[:-2]
                            pred_row = df_gantt[df_gantt['ID_WBS'].astype(str) == p_id]
                            if not pred_row.empty:
                                pred_y = p_id + " - " + str(pred_row.iloc[0]['Attività'])
                                if vista == "Esecuzione (Esecutivo)" and pd.notna(pred_row.iloc[0]['Fine_Effettiva']):
                                    pred_end = pred_row.iloc[0]['Fine_Effettiva'] + pd.Timedelta(days=1)
                                else:
                                    pred_end = pred_row.iloc[0]['Fine_Prevista'] + pd.Timedelta(days=1)
                                
                                is_critical = cpm_data.get(wbs_id, {}).get('is_critical', False)
                                pred_is_critical = cpm_data.get(p_id, {}).get('is_critical', False)
                                arrow_color = '#D32F2F' if (is_critical and pred_is_critical) else '#9E9E9E'
                                fig.add_annotation(
                                    x=succ_start, y=succ_y, ax=pred_end, ay=pred_y,
                                    xref='x', yref='y', axref='x', ayref='y',
                                    showarrow=True, arrowhead=2, arrowsize=1, arrowwidth=1.5,
                                    arrowcolor=arrow_color, opacity=0.8, standoff=2, startstandoff=2
                                )
            
            altezza_dinamica = max(500, len(df_gantt) * 45)
            fig.update_layout(
                barmode='overlay', height=altezza_dinamica, bargap=0.3, 
                xaxis_title="Linea Temporale", yaxis_title="Lavorazioni (WBS)", 
                yaxis={'autorange': 'reversed'}, xaxis_type='date',
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor="rgba(255,255,255,0.8)")
            )
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("⚠️ Il cronoprogramma è vuoto. Inserisci le date di Inizio e Fine nel Tab 1.")

        # SCUDO ANTI-CRASH E RETRO-COMPATIBILITÀ ---
        # Garantisce che le colonne esistano sempre, anche su vecchi salvataggi
        if 'Vincolo_Burocratico' not in st.session_state.wbs_data.columns:
            st.session_state.wbs_data['Vincolo_Burocratico'] = 'Nessuno'
        if 'Vincolo_Assolto' not in st.session_state.wbs_data.columns:
            st.session_state.wbs_data['Vincolo_Assolto'] = False

        # ========================================================
        # 2. SCADENZARIO AMMINISTRATIVO
        # ========================================================
        st.subheader("🏛️ Scadenzario Autorizzazioni (Cancelli Burocratici)")
        st.markdown("Monitora le lavorazioni che necessitano di un'autorizzazione formale prima di poter iniziare. Le voci già assolte rimangono in archivio per consultazione.")
        
        # FIX CHIRURGICO 2: Prendiamo TUTTO l'albero WBS, non solo le foglie!
        df_wbs_scad = st.session_state.wbs_data.copy()
        
        # FIX CHIRURGICO 3: Pulizia del testo a prova di bomba per catturare "nessuno", "Nessuno", o celle vuote
        df_wbs_scad['Vincolo_Clean'] = df_wbs_scad['Vincolo_Burocratico'].astype(str).str.strip().str.lower()
        df_vincoli = df_wbs_scad[~df_wbs_scad['Vincolo_Clean'].isin(['nessuno', 'nan', 'none', ''])].copy()
        
        if not df_vincoli.empty:
            df_vincoli['Assolto'] = df_vincoli['Vincolo_Assolto'].apply(lambda x: True if str(x).strip().lower() in ['true', '1', 't', 'y', 'yes'] else False)
            
            oggi_scad = pd.Timestamp.today().date()
            dati_scadenzario = []
            
            for _, row in df_vincoli.iterrows():
                inizio_prev = pd.to_datetime(row['Inizio_Previsto'], errors='coerce')
                assolto = row['Assolto']
                
                if assolto:
                    stato = "✅ Assolto (Autorizzato)"
                    sort_val = 9999  # Spinge in fondo
                    data_str = inizio_prev.strftime('%d/%m/%Y') if pd.notna(inizio_prev) else "N/D"
                else:
                    if pd.notna(inizio_prev):
                        giorni_mancanti = (inizio_prev.date() - oggi_scad).days
                        if giorni_mancanti < 0:
                            stato = f"🔴 SCADUTO (Ritardo: {abs(giorni_mancanti)} gg)"
                        elif giorni_mancanti <= 15:
                            stato = f"🟠 CRITICO (Scade tra {giorni_mancanti} gg)"
                        else:
                            stato = f"🟢 In tempo ({giorni_mancanti} gg)"
                        sort_val = giorni_mancanti
                        data_str = inizio_prev.strftime('%d/%m/%Y')
                    else:
                        stato = "⚪ Data Inizio non definita"
                        sort_val = 999
                        data_str = "N/D"
                    
                dati_scadenzario.append({
                    'WBS': f"{row['ID_WBS']} - {row['Attività']}",
                    'Autorizzazione Richiesta': row['Vincolo_Burocratico'],
                    'Scadenza (Inizio Previsto)': data_str,
                    'Stato Urgenza': stato,
                    '_sort': sort_val
                })
                
            df_display_scad = pd.DataFrame(dati_scadenzario).sort_values('_sort').drop(columns=['_sort'])
            
            def colora_stato(val):
                if '🔴' in str(val): return 'color: white; background-color: #D32F2F; font-weight: bold;'
                if '🟠' in str(val): return 'color: white; background-color: #FF9800; font-weight: bold;'
                if '🟢' in str(val): return 'color: white; background-color: #388E3C; font-weight: bold;'
                if '✅' in str(val): return 'color: #155724; background-color: #C8E6C9; font-weight: bold;' 
                return ''
                
            st.dataframe(df_display_scad.style.map(colora_stato, subset=['Stato Urgenza']), use_container_width=True, hide_index=True)
                
        else:
            st.info("ℹ️ Nessun vincolo burocratico attualmente impostato nel Tab 1.")
    
    #-------------------------------    
    # --- TAB 5: EVM E CASH FLOW ---
    #-------------------------------
    with tab5:
        st.header("Controllo Costi e Analisi EVM")

        # ======================================================
        # ALLERTA FINANZIARIA GIANFRY: VARIANTI DA QUANTIFICARE
        if 'tickets_data' in st.session_state and not st.session_state.tickets_data.empty:
            df_t_alert = st.session_state.tickets_data
            if 'Variazione_Costi' in df_t_alert.columns and 'Variazione_Tempi' in df_t_alert.columns:
                
                # Conversione sicura per scovare le celle "" (vuote)
                costi_vuoti = pd.to_numeric(df_t_alert['Variazione_Costi'], errors='coerce').isna()
                tempi_vuoti = pd.to_numeric(df_t_alert['Variazione_Tempi'], errors='coerce').isna()
                
                # LOGICA CORRETTA: Suona solo se sono vuoti ENTRAMBI
                da_compilare = df_t_alert[
                    (df_t_alert['Stato'] == 'Approvato ✅') & 
                    (df_t_alert['Tipologia'] == 'Richiesta di Variante') & 
                    (costi_vuoti & tempi_vuoti)
                ]
                
                if not da_compilare.empty:
                    wbs_sospese = da_compilare['ID_WBS_Rif'].astype(str).tolist()
                    st.error(f"🚨 **ALLERTA FINANZIARIA (GIANFRY):** Hai approvato una Variante per le WBS **{', '.join(wbs_sospese)}** senza quantificarne l'impatto! Vai nel Tab 9 e compila 'Variazione Costi' o 'Variazione Tempi' per ricalibrare il motore EVM.")
        # ======================================================
        
        data_status_evm = st.date_input("📅 Data di Stato (Status Date):", value=pd.Timestamp.today().date())
        df_evm = calcola_evm(get_foglie(st.session_state.wbs_data), data_status_evm)       
        tot_bac, tot_pv, tot_ev, tot_ac = df_evm['BAC_Budget'].sum(), df_evm['PV'].sum(), df_evm['EV'].sum(), df_evm['AC_Costo_Reale'].sum()
        tot_eac, tot_etc, tot_vac = df_evm['EAC'].sum(), df_evm['ETC'].sum(), df_evm['VAC'].sum()
        cpi_globale = tot_ev / tot_ac if tot_ac > 0 else 1.0
        spi_globale = tot_ev / tot_pv if tot_pv > 0 else 1.0
        perc_completamento = (tot_ev / tot_bac * 100) if tot_bac > 0 else 0.0
        perc_pianificata = (tot_pv / tot_bac * 100) if tot_bac > 0 else 0.0
        
        # --- CALCOLO CONTINGENCY RESERVE (EMV DAI RISCHI) ---
        df_rischi = st.session_state.rischi_data.copy()
        contingency_reserve = 0.0
        
        if not df_rischi.empty:
            # Peschiamo SOLO i rischi attivi o monitorati
            rischi_attivi = df_rischi[df_rischi['Stato'].isin(['Attivo ▾', 'Monitorato ▾'])]
            for _, r_row in rischi_attivi.iterrows():
                wbs_id_full = str(r_row.get('ID_WBS_Rif', ''))
                wbs_id = wbs_id_full.split(' - ')[0].strip()
                
                # Cerchiamo il budget di quella specifica attività
                wbs_match = df_evm[df_evm['ID_WBS'].astype(str) == wbs_id]
                wbs_budget = float(wbs_match.iloc[0]['BAC_Budget']) if not wbs_match.empty else 0.0
                
                prob = pd.to_numeric(r_row['Probabilità (1-5)'], errors='coerce')
                imp = pd.to_numeric(r_row['Impatto (1-5)'], errors='coerce')
                
                if pd.notna(prob) and pd.notna(imp) and wbs_budget > 0:
                    prob_pct = prob * 0.20        # Prob: 1=20%, 5=100%
                    imp_pct = (imp - 1) * 0.05    # Impatto: 1=0%, 5=20% del budget
                    
                    emv = wbs_budget * prob_pct * imp_pct
                    contingency_reserve += emv
                    
        eac_risk_adjusted = tot_eac + contingency_reserve
        
        # --- CRUSCOTTO UI ---
        col_box1, col_box2 = st.columns(2)
        with col_box1:
            with st.container(border=True):
                st.markdown("#### 📊 Stato Attuale (Consuntivo)")
                c1, c2, c3 = st.columns(3)
                c1.metric("Budget Totale (BAC)", f"€ {tot_bac:,.0f}")
                c2.metric("Lavoro Eseguito (EV)", f"€ {tot_ev:,.0f}")
                c3.metric("Costi Sostenuti (AC)", f"€ {tot_ac:,.0f}")
                st.divider()
                c4, c5 = st.columns(2)
                c4.metric("Avanzamento Globale", f"{perc_completamento:.1f}%", delta=f"Pianificato: {perc_pianificata:.1f}%", delta_color="off")
                c5.metric("SPI (Tempi)", f"{spi_globale:.2f}", delta="In ritardo" if spi_globale < 1 else "In anticipo", delta_color="inverse")

        with col_box2:
            with st.container(border=True):
                st.markdown("#### 🔮 Previsioni a Finire (Proiezioni)")
                c1, c2, c3 = st.columns(3)
                c1.metric("Costo Stimato (EAC)", f"€ {tot_eac:,.0f}", delta="Puro EVM matematico", delta_color="off")
                c2.metric("Costo Residuo (ETC)", f"€ {tot_etc:,.0f}", delta="Capitale necessario", delta_color="off")
                c3.metric("Varianza a Finire (VAC)", f"€ {tot_vac:,.0f}", delta="Perdita" if tot_vac < 0 else "Risparmio", delta_color="normal")
                st.divider()
                c4, c5 = st.columns(2)
                c4.metric("CPI (Costi)", f"{cpi_globale:.2f}", delta="Over-budget" if cpi_globale < 1 else "Under-budget", delta_color="inverse")
        
        # ========================================================
        # CRUSCOTTO CASH FLOW NETTO
        # ========================================================
        st.divider()
        st.markdown("### 💶 Esposizione Finanziaria - Cash Flow Netto")
        
        # Calcolo sicuro delle uscite totali leggendo direttamente la colonna Actual Cost (AC)
        tot_uscite_ac = df_evm['AC_Costo_Reale'].sum() if not df_evm.empty else 0.0
        
        # Sommiamo le entrate SOLO se i SAL sono stati effettivamente incassati
        df_sal_calcolo = st.session_state.sal_data
        tot_entrate_sal = 0.0
        if not df_sal_calcolo.empty:
            sal_pagati = df_sal_calcolo[df_sal_calcolo['Stato_Pagamento'] == 'Pagato ▾']
            tot_entrate_sal = pd.to_numeric(sal_pagati['Importo_Euro'], errors='coerce').fillna(0).sum()
            
        cash_flow_netto = tot_entrate_sal - tot_uscite_ac
        
        c_cf1, c_cf2, c_cf3 = st.columns(3)
        c_cf1.metric("Totale Uscite Sostenute (AC)", f"€ {tot_uscite_ac:,.0f}")
        c_cf2.metric("Totale Incassi (SAL Pagati)", f"€ {tot_entrate_sal:,.0f}")
        c_cf3.metric("Cash Flow Netto (Liquidità)", f"€ {cash_flow_netto:,.0f}", 
                     delta="Esposizione Negativa!" if cash_flow_netto < 0 else "Liquidità Positiva", 
                     delta_color="normal" if cash_flow_netto >= 0 else "inverse")
                     
        if cash_flow_netto < 0:
            st.error(f"⚠️ **ALLERTA LIQUIDITÀ:** L'esposizione finanziaria del cantiere è negativa per **€ {abs(cash_flow_netto):,.0f}**. Stai anticipando capitale rispetto a quanto incassato. Valuta l'emissione immediata di un nuovo SAL per ripristinare il flusso di cassa.")
        elif cash_flow_netto > 0:
            st.success(f"✅ **LIQUIDITÀ ATTIVA:** Il cantiere si sta autofinanziando correttamente. Hai un margine di cassa positivo di **€ {cash_flow_netto:,.0f}**.")
        
        # --- SCUDO FINANZIARIO  ---

        st.divider()
        # 1. Motore di calcolo EMV (Expected Monetary Value)
        df_rischi = st.session_state.rischi_data.copy()
        fondo_imprevisti = 0.0
        conteggio_attivi = 0
        
        if not df_rischi.empty:
            # Peschiamo SOLO i rischi non ancora risolti
            rischi_attivi = df_rischi[df_rischi['Stato'].isin(['Attivo ▾', 'Monitorato ▾'])]
            conteggio_attivi = len(rischi_attivi)
            
            if conteggio_attivi > 0:
                # Mappatura standard Punteggio -> Percentuale
                prob_map = {1: 0.10, 2: 0.30, 3: 0.50, 4: 0.70, 5: 0.90} # 10% - 90%
                imp_map = {1: 0.02, 2: 0.05, 3: 0.10, 4: 0.15, 5: 0.20} # 2% - 20% del Budget
                
                for _, r in rischi_attivi.iterrows():
                    id_wbs = str(r['ID_WBS_Rif']).split(' - ')[0].strip()
                    try:
                        prob = int(r['Probabilità (1-5)'])
                        imp = int(r['Impatto (1-5)'])
                        
                        if prob in prob_map and imp in imp_map:
                            # Troviamo il Budget originale della singola WBS interessata
                            wbs_row = df_evm[df_evm['ID_WBS'].astype(str) == id_wbs]
                            bac_wbs = wbs_row['BAC_Budget'].values[0] if not wbs_row.empty else 0.0
                            
                            # Calcolo EMV per questo specifico rischio
                            emv = bac_wbs * prob_map[prob] * imp_map[imp]
                            fondo_imprevisti += emv
                    except:
                        pass
        
        eac_risk_adjusted = tot_eac + fondo_imprevisti

        # 2. Interfaccia Visiva dello Scudo
        st.markdown("### 🛡️ Scudo Finanziario (Risk-Adjusted EVM)")
        with st.expander("Metodologia Contingency Reserve (Fondo Imprevisti)", expanded=True):
            st.markdown(f'''
            **Integrazione EMV (Expected Monetary Value):**
            Il sistema non si limita a calcolare la proiezione dei costi attuale (EAC Tradizionale), ma "congela" dinamicamente una quota di capitale in base ai pericoli registrati nel **Risk Register**. 
            Per ogni rischio ancora *Attivo* o *Monitorato*, il motore trasforma Probabilità e Impatto in percentuali, calcolando la penale attesa sul Budget (BAC) della specifica lavorazione.
            
            * **Formula EMV Singolo:** `BAC Lavorazione` × `Probabilità (%)` × `Impatto (%)`
            * **Formula EAC Risk-Adjusted:** `EAC Tradizionale` + `Σ (Somma EMV Attivi)`
            
            *(N.B. Appena modifichi lo stato di un rischio in "Mitigato" o "Chiuso" nel Tab 8, il Fondo si svuoterà automaticamente, liberando quel capitale per il committente).*
            ''')
            
            c_scudo1, c_scudo2, c_scudo3 = st.columns(3)
            c_scudo1.metric("EAC Tradizionale (Puro)", f"€ {tot_eac:,.0f}", help="Costo stimato a fine progetto senza considerare i rischi futuri.")
            c_scudo2.metric("Fondo Imprevisti (EMV Totale)", f"€ {fondo_imprevisti:,.0f}", delta=f"{conteggio_attivi} Rischi aperti", delta_color="inverse")
            c_scudo3.metric("EAC Risk-Adjusted", f"€ {eac_risk_adjusted:,.0f}", delta="Peggior Scenario Probabile", delta_color="off")
        
        st.divider()

        st.subheader("📈 Andamento di Progetto & Proiezioni")
        
        df_scurve = genera_dati_scurve(df_evm, st.session_state.registro_data, data_status_evm)
        if df_scurve is not None and not df_scurve.empty:
            fig_scurve = px.line(df_scurve, x='Data', y=['PV (Valore Pianificato)', 'EV (Valore Guadagnato)', 'AC (Costo Reale)'], color_discrete_map={'PV (Valore Pianificato)': 'blue', 'EV (Valore Guadagnato)': 'green', 'AC (Costo Reale)': 'red'}, labels={'value': 'Importo (€)', 'variable': 'Metrica EVM'})
            df_past = df_scurve[df_scurve['Data'] <= data_status_evm]
            if not df_past.empty:
                min_date, max_date = df_scurve['Data'].min(), df_scurve['Data'].max()
                spi_effettivo = df_past.iloc[-1]['EV (Valore Guadagnato)'] / df_past.iloc[-1]['PV (Valore Pianificato)'] if df_past.iloc[-1]['PV (Valore Pianificato)'] > 0 else 1.0
                giorni_stimati = min(int((max_date - min_date).days / spi_effettivo) if spi_effettivo > 0 else (max_date - min_date).days, (max_date - min_date).days * 3) 
                data_fine_stimata = min_date + pd.Timedelta(days=giorni_stimati)
                
                # Proiezione Base
                fig_scurve.add_trace(go.Scatter(x=[data_status_evm, data_fine_stimata], y=[df_past.iloc[-1]['AC (Costo Reale)'], tot_eac], mode='lines', line=dict(color='red', dash='dot', width=2), name='Proiezione Costi (EVM Puro)'))
                # Proiezione Risk-Adjusted (se ci sono rischi)
                if contingency_reserve > 0:
                    fig_scurve.add_trace(go.Scatter(x=[data_status_evm, data_fine_stimata], y=[df_past.iloc[-1]['AC (Costo Reale)'], eac_risk_adjusted], mode='lines', line=dict(color='darkred', dash='solid', width=3), name='Proiezione Risk-Adjusted'))
                
                fig_scurve.add_trace(go.Scatter(x=[data_status_evm, data_fine_stimata], y=[df_past.iloc[-1]['EV (Valore Guadagnato)'], tot_bac], mode='lines', line=dict(color='green', dash='dot', width=2), name='Proiezione Lavoro (Tempi)'))
                fig_scurve.update_xaxes(range=[min_date, max(max_date, data_fine_stimata) + pd.Timedelta(days=5)])
            fig_scurve.update_layout(hovermode="x unified", legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1), margin=dict(t=50, b=20))
            st.plotly_chart(fig_scurve, use_container_width=True)
        else:
            st.info("ℹ️ Date insufficienti per la Curva ad S.")
        
        st.divider()
        st.subheader("Raffronto Costi per Attività")
        if not df_evm.empty:
            fig_evm = go.Figure(data=[
                go.Bar(name='BAC', x=df_evm['Attività'], y=df_evm['BAC_Budget'], marker_color='lightgray', text=df_evm['BAC_Budget'], texttemplate='€ %{text:,.0f}', textposition='outside', textangle=-90),
                go.Bar(name='EV', x=df_evm['Attività'], y=df_evm['EV'], marker_color='green', text=df_evm['EV'], texttemplate='€ %{text:,.0f}', textposition='outside', textangle=-90),
                go.Bar(name='AC', x=df_evm['Attività'], y=df_evm['AC_Costo_Reale'], marker_color='red', text=df_evm['AC_Costo_Reale'], texttemplate='€ %{text:,.0f}', textposition='outside', textangle=-90)
            ])
            fig_evm.update_layout(barmode='group', margin=dict(t=80), uniformtext_minsize=9, uniformtext_mode='hide')
            st.plotly_chart(fig_evm, use_container_width=True)
            
        col_KPI, col_LEGENDA = st.columns([7, 3]) 
        with col_KPI:
            st.subheader("Indicatori di Performance (KPI)")
            df_kpi = df_evm[['Attività', '%_Completamento', 'CPI', 'SPI', 'CV']].copy()
            def color_kpi(val):
                if isinstance(val, (int, float)):
                    if val < 0.95: return 'color: red; font-weight: bold;'
                    elif val >= 1.0: return 'color: green'
                return ''
            if not df_kpi.empty:
                st.dataframe(df_kpi.style.map(color_kpi, subset=['CPI', 'SPI']).format({'CPI': "{:.2f}", 'SPI': "{:.2f}", 'CV': "€ {:.2f}"}), use_container_width=True)

        with col_LEGENDA:
            st.subheader("Legenda EVM")
            st.markdown("* **CPI:** Efficienza costi (<1 sforamento budget)\n* **SPI:** Efficienza tempi (<1 in ritardo)\n* **CV:** Varianza Costi Assoluta")
            
    # --- TAB 6: GESTIONE FINANZIARIA E CASH FLOW ---
    with tab6:
        st.header("Gestione Finanziaria e Liquidità")
        
        tab_uscite, tab_entrate = st.tabs(["💸 Uscite (Costi Reali - AC)", "💰 Entrate (SAL Certificati)"])
        
        with tab_uscite:
            st.subheader("Brogliaccio Spese e Fatture Fornitori")
            st.markdown("Le spese registrate qui alimentano il Costo Reale (AC) delle singole WBS.")
            
            df_reg = st.session_state.registro_data.copy()
            if not df_reg.empty:
                df_reg['Data'] = pd.to_datetime(df_reg['Data'], errors='coerce').dt.date
            else:
                df_reg['Data'] = pd.Series(dtype='object')
            st.session_state.registro_data = df_reg
            
            leaf_wbs_reg = get_foglie(st.session_state.wbs_data)
            wbs_options = [f"{row['ID_WBS']} - {row['Attività']}" for _, row in leaf_wbs_reg.iterrows()]
            
            edited_registro = st.data_editor(
            st.session_state.registro_data, num_rows="dynamic", use_container_width=True, hide_index=True,
            column_config={
                "Data": st.column_config.DateColumn("Data Fattura / Spesa"),
                "Voce_WBS": st.column_config.SelectboxColumn("Attività WBS (Rif.)", options=wbs_options), # NOME CORRETTO
                "Descrizione": st.column_config.TextColumn("Descrizione", width="medium"), # NOME CORRETTO
                "Importo_Netto": st.column_config.NumberColumn("Importo (€)", format="€ %.2f") # NOME CORRETTO
            }
        )
            if st.button("💾 SALVA REGISTRO USCITE", type="primary", use_container_width=True):
                st.session_state.registro_data = edited_registro
                st.success("✅ Registro Spese salvato e costi riallineati!")
                st.rerun()

        with tab_entrate:
            st.subheader("Emissione SAL e Incassi Committenza")
            st.markdown("Registra gli Stati Avanzamento Lavori emessi e pagati. Solo i SAL in stato **'Pagato'** concorreranno al calcolo della liquidità attiva (Cash Flow).")
            
            df_sal = st.session_state.sal_data.copy()
            if not df_sal.empty:
                df_sal['Data_Emissione'] = pd.to_datetime(df_sal['Data_Emissione'], errors='coerce').dt.date
            else:
                df_sal['Data_Emissione'] = pd.Series(dtype='object')
            
            edited_sal = st.data_editor(
                df_sal, num_rows="dynamic", use_container_width=True, hide_index=True,
                column_config={
                    "Data_Emissione": st.column_config.DateColumn("Data SAL"),
                    "Descrizione_SAL": st.column_config.TextColumn("Riferimento / Note", width="medium"),
                    "Importo_Euro": st.column_config.NumberColumn("Importo (€)", format="€ %.2f", min_value=0.0),
                    "Stato_Pagamento": st.column_config.SelectboxColumn("Stato", options=["Emesso ▾", "Pagato ▾"])
                }
            )
            if st.button("💾 SALVA REGISTRO ENTRATE (SAL)", type="primary", use_container_width=True):
                st.session_state.sal_data = edited_sal
                st.success("✅ Registro Incassi aggiornato con successo!")
                st.rerun()

        # ========================================================
        # GRAFICO: ANDAMENTO FLUSSI DI CASSA (CUMULATIVO)
        # ========================================================
        st.divider()
        st.subheader("📊 Andamento Flussi di Cassa (Cash Flow Storico)")
        
        df_grafico_uscite = st.session_state.registro_data.copy()
        df_grafico_entrate = st.session_state.sal_data.copy()
        
        oggi = pd.Timestamp.today().normalize()
        
        # 1. Preparazione Dati Uscite (Spese Cumulative) - ORA RICONOSCE "Importo_Netto"
        if not df_grafico_uscite.empty:
            # Creiamo una lista di tutti i possibili nomi che hai usato per i costi
            possibili_nomi_costo = ['Importo_Netto', 'Importo_Euro', 'Costo', 'Importo']
            col_imp_uscite = next((col for col in possibili_nomi_costo if col in df_grafico_uscite.columns), None)
            
            # Stessa cosa per la data
            possibili_nomi_data = ['Data', 'Data Fattura / Spesa', 'Data_Fattura']
            col_data_uscite = next((col for col in possibili_nomi_data if col in df_grafico_uscite.columns), None)
            
            if col_data_uscite and col_imp_uscite:
                df_grafico_uscite['Data_Grafico'] = pd.to_datetime(df_grafico_uscite[col_data_uscite], errors='coerce')
                df_grafico_uscite[col_imp_uscite] = pd.to_numeric(df_grafico_uscite[col_imp_uscite], errors='coerce').fillna(0)
                df_grafico_uscite = df_grafico_uscite.dropna(subset=['Data_Grafico'])
                df_grafico_uscite = df_grafico_uscite.sort_values('Data_Grafico')
                df_grafico_uscite['Cumulato'] = df_grafico_uscite[col_imp_uscite].cumsum()
            else:
                df_grafico_uscite = pd.DataFrame()
            
        # 2. Preparazione Dati Entrate (Solo SAL Pagati, Cumulativi)
        if not df_grafico_entrate.empty and 'Stato_Pagamento' in df_grafico_entrate.columns:
            df_grafico_entrate = df_grafico_entrate[df_grafico_entrate['Stato_Pagamento'] == 'Pagato ▾'].copy()
            col_imp_entrate = 'Importo_Euro' if 'Importo_Euro' in df_grafico_entrate.columns else 'Importo'
            
            if 'Data_Emissione' in df_grafico_entrate.columns and col_imp_entrate in df_grafico_entrate.columns:
                df_grafico_entrate['Data_Grafico'] = pd.to_datetime(df_grafico_entrate['Data_Emissione'], errors='coerce')
                df_grafico_entrate[col_imp_entrate] = pd.to_numeric(df_grafico_entrate[col_imp_entrate], errors='coerce').fillna(0)
                df_grafico_entrate = df_grafico_entrate.dropna(subset=['Data_Grafico'])
                df_grafico_entrate = df_grafico_entrate.sort_values('Data_Grafico')
                df_grafico_entrate['Cumulato'] = df_grafico_entrate[col_imp_entrate].cumsum()
            else:
                df_grafico_entrate = pd.DataFrame()

        # 3. Disegno del Grafico
        if df_grafico_uscite.empty and df_grafico_entrate.empty:
            st.info("Aggiungi spese o incassi (e ricordati di cliccare SALVA) per visualizzare il grafico.")
        else:
            fig_cf = go.Figure()
            
            # Traccia Costi (Rossa)
            if not df_grafico_uscite.empty:
                # Estensione linea fino a oggi per continuità visiva
                ultima_data_u = df_grafico_uscite['Data_Grafico'].max()
                if ultima_data_u < oggi:
                    nuova_riga = pd.DataFrame({'Data_Grafico': [oggi], 'Cumulato': [df_grafico_uscite['Cumulato'].iloc[-1]]})
                    df_grafico_uscite = pd.concat([df_grafico_uscite, nuova_riga], ignore_index=True)
                    
                fig_cf.add_trace(go.Scatter(
                    x=df_grafico_uscite['Data_Grafico'],
                    y=df_grafico_uscite['Cumulato'],
                    mode='lines+markers',
                    name='Uscite Cumulate (AC)',
                    line=dict(color='#D32F2F', width=3, shape='hv'), # 'hv' = a gradini
                    fill='tozeroy',
                    fillcolor='rgba(211, 47, 47, 0.1)'
                ))
                
            # Traccia Incassi (Verde)
            if not df_grafico_entrate.empty:
                # Estensione linea fino a oggi per continuità visiva
                ultima_data_e = df_grafico_entrate['Data_Grafico'].max()
                if ultima_data_e < oggi:
                    nuova_riga = pd.DataFrame({'Data_Grafico': [oggi], 'Cumulato': [df_grafico_entrate['Cumulato'].iloc[-1]]})
                    df_grafico_entrate = pd.concat([df_grafico_entrate, nuova_riga], ignore_index=True)

                fig_cf.add_trace(go.Scatter(
                    x=df_grafico_entrate['Data_Grafico'],
                    y=df_grafico_entrate['Cumulato'],
                    mode='lines+markers',
                    name='Incassi Cumulati (SAL)',
                    line=dict(color='#4CAF50', width=3, shape='hv'), # 'hv' = a gradini
                    fill='tozeroy',
                    fillcolor='rgba(76, 175, 80, 0.1)'
                ))
                
            fig_cf.update_layout(
                xaxis_title="Linea Temporale",
                yaxis_title="Importo Cumulato (€)",
                hovermode="x unified",
                height=400,
                legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1, bgcolor="rgba(255,255,255,0.8)")
            )
            
            st.plotly_chart(fig_cf, use_container_width=True)
            
    # --- TAB 7: DIREZIONE LAVORI, CAPA & REPORTISTICA ---
    with tab7:
        st.header("Direzione Lavori: Interventi (CAPA) e Simulazioni")
        
        leaf_wbs_capa = get_foglie(st.session_state.wbs_data)
        wbs_options_capa = [f"{row['ID_WBS']} - {row['Attività']}" for _, row in leaf_wbs_capa.iterrows()]
        obs_options_capa = [f"{row['ID_OBS']} - {row['Risorsa']}" for _, row in st.session_state.obs_data.iterrows()]
        
        # NUOVO: Recuperiamo i rischi registrati dal Tab 8 per popolare la tendina
        lista_rischi = [""] + [str(r) for r in st.session_state.rischi_data['Descrizione_Rischio'].dropna().unique() if str(r).strip() != ""]
        
        st.subheader("1. Registro Azioni Correttive e Preventive (CAPA)")
        
        df_capa = st.session_state.capa_data.copy()
        
        # Retro-compatibilità: Se carichi un vecchio progetto senza questa colonna, la crea al volo
        if 'Rischio_Associato' not in df_capa.columns:
            df_capa['Rischio_Associato'] = ""
            
        if not df_capa.empty:
            df_capa['Data_Apertura'] = pd.to_datetime(df_capa['Data_Apertura'], errors='coerce').dt.date
        else:
            df_capa['Data_Apertura'] = pd.Series(dtype='object')
        st.session_state.capa_data = df_capa
        
        edited_capa = st.data_editor(
            st.session_state.capa_data, num_rows="dynamic", use_container_width=True, hide_index=True,
            column_config={
                "Data_Apertura": st.column_config.DateColumn("Data Segnalazione"),
                "ID_WBS_Rif": st.column_config.SelectboxColumn("Attività WBS (Rif.)", options=wbs_options_capa),
                "Tipo_Azione": st.column_config.SelectboxColumn("Tipo", options=["Correttiva ▾", "Preventiva ▾"]),
                "Descrizione": st.column_config.TextColumn("Descrizione Intervento / Ordine", width="medium"),
                "Responsabile_OBS": st.column_config.SelectboxColumn("Risorsa (OBS)", options=obs_options_capa),
                "Rischio_Associato": st.column_config.SelectboxColumn("Rischio da mitigare ▾", options=lista_rischi), 
                "Costo_Intervento": st.column_config.NumberColumn("Costo Intervento (€)", format="€ %.2f", default=0.0, help="Costo della non-qualità"),
                "Giorni_Intervento": st.column_config.NumberColumn("Giorni extra", step=1, default=0),
                "Costo_Scaricato": None, # Nasconde la colonna di sistema
                "Stato": st.column_config.SelectboxColumn("Stato", options=["Aperto ▾", "In Lavorazione ▾", "Chiuso ▾"])
            }
        )
        
        st.divider()
        st.warning("⚠️ **Ricordati di cliccare il tasto rosso qui sotto dopo aver inserito i dati!**")
        if st.button("💾 SALVA REGISTRO CAPA", type="primary", use_container_width=True):
            df_capa_new = edited_capa.copy()
            
            # Variabili per i messaggi a schermo
            rischi_aggiornati = False
            costi_inviati = 0
            df_rischi = st.session_state.rischi_data.copy()
            nuove_spese = []
            
            for idx, capa_row in df_capa_new.iterrows():
                # A) AUTO-MITIGAZIONE RISCHI (IL PILOTA AUTOMATICO)
                if capa_row['Stato'] == 'Chiuso ▾' and pd.notna(capa_row.get('Rischio_Associato')) and str(capa_row.get('Rischio_Associato')).strip() != "":
                    rischio_target = str(capa_row['Rischio_Associato']).strip()
                    mask = (df_rischi['Descrizione_Rischio'] == rischio_target) & (~df_rischi['Stato'].isin(['Mitigato ▾', 'Chiuso ▾']))
                    if mask.any():
                        df_rischi.loc[mask, 'Probabilità (1-5)'] = 1
                        df_rischi.loc[mask, 'Stato'] = 'Mitigato ▾'
                        rischi_aggiornati = True

               # B) INVIO AUTOMATICO COSTI DELLA NON-QUALITÀ AL TAB 6
                val_scaricato = capa_row.get('Costo_Scaricato', False)
                # Decodificatore a prova di bomba per il JSON
                gia_scaricato = val_scaricato if isinstance(val_scaricato, bool) else str(val_scaricato).strip().lower() in ['true', '1', 't', 'y', 'yes']
                
                if capa_row['Stato'] == 'Chiuso ▾' and not gia_scaricato:
                    costo = pd.to_numeric(capa_row.get('Costo_Intervento', 0), errors='coerce')
                    if pd.notna(costo) and costo > 0:
                        nuova_spesa = {
                            'Data': pd.Timestamp.today().date(),
                            'N_Doc': f"CAPA-{idx}",
                            'Fornitore': str(capa_row.get('Responsabile_OBS', '')),
                            'Voce_WBS': str(capa_row.get('ID_WBS_Rif', '')),
                            'Importo_Netto': costo,
                            'Descrizione': f"🔴 COSTO NON-QUALITÀ: {str(capa_row.get('Descrizione', ''))}"
                        }
                        nuove_spese.append(nuova_spesa)
                        df_capa_new.at[idx, 'Costo_Scaricato'] = True
                        costi_inviati += 1
            
            # Applica le modifiche ai database
            if nuove_spese:
                st.session_state.registro_data = pd.concat([st.session_state.registro_data, pd.DataFrame(nuove_spese)], ignore_index=True)
            
            if rischi_aggiornati:
                st.session_state.rischi_data = df_rischi
                
            st.session_state.capa_data = df_capa_new
            
            # Feedback intelligente a schermo
            if costi_inviati > 0:
                st.success(f"✅ Interventi salvati! Sono stati inviati **{costi_inviati} Costi di Non-Qualità** al Registro Contabile (Tab 6).")
            elif rischi_aggiornati:
                st.success("✅ Interventi salvati! Il pilota automatico ha **Mitigato** i rischi associati alle azioni chiuse nel Tab 8.")
            else:
                st.success("✅ Interventi salvati con successo nel database!")

            aggiorna_costi_reali()
            st.session_state.wbs_data = aggiorna_gerarchia(st.session_state.wbs_data)
            for k in list(st.session_state.keys()):
                if k.startswith("editor_wbs_"):
                    del st.session_state[k]
            
            import time
            time.sleep(2)
            st.rerun()
            
        # ------------------------------
        # SIMULAZIONE
        # ------------------------------
        
        with st.expander("🔬 2. Ambiente di Simulazione (Compromesso Costi / Tempi)"):
            st.markdown("Simula l'impatto di un'azione correttiva senza alterare i principi base dell'EVM. Valuta costi di *Crashing* e l'incidenza sui Costi Indiretti di cantiere.")
            
            c_sim1, c_sim2, c_sim3, c_sim4 = st.columns([2, 1.5, 1.5, 1.5])
            wp_scelto = c_sim1.selectbox("Seleziona Work Package da simulare", options=wbs_options_capa)
            extra_costo = c_sim2.number_input("Costo Extra Diretto (€)", value=0.0, step=500.0, help="Costo vivo aggiuntivo (es. premi, straordinari). Intacca il CPI e fa salire l'EAC.")
            var_giorni = c_sim3.number_input("Variazione Tempi (Giorni)", value=0, step=1, help="Usa numeri NEGATIVI per anticipare (es. -5), POSITIVI per ritardare (es. 5).")
            costo_indiretto_gg = c_sim4.number_input("Costi Indiretti (€/gg)", value=0.0, step=50.0, help="Costi fissi per ogni giorno di ritardo (es. gru, baraccamenti, penali)")
            
            if wp_scelto:
                wp_id = wp_scelto.split(' - ')[0]
                
                # --- CALCOLO REALE (BASE) ---
                oggi = pd.Timestamp.today().date()
                df_reale_calc = calcola_evm(get_foglie(st.session_state.wbs_data), oggi)
                
                eac_attuale = df_reale_calc['EAC'].sum()
                ac_attuale_tot = df_reale_calc['AC_Costo_Reale'].sum()
                
                # --- SIMULAZIONE EVM PURA ---
                df_simulazione = st.session_state.wbs_data.copy()
                indice_riga = df_simulazione.index[df_simulazione['ID_WBS'] == wp_id].tolist()
                if indice_riga:
                    idx = indice_riga[0]
                    # Iniettare costo diretto abbassa il CPI e alza l'EAC secondo le regole pure dell'EVM
                    df_simulazione.at[idx, 'AC_Costo_Reale'] = pd.to_numeric(df_simulazione.at[idx, 'AC_Costo_Reale'], errors='coerce') + extra_costo
                    
                df_sim_calc = calcola_evm(get_foglie(df_simulazione), oggi)
                eac_simulato = df_sim_calc['EAC'].sum()
                ac_simulato_tot = df_sim_calc['AC_Costo_Reale'].sum()
                
                # --- SIMULAZIONE TEMPI E COSTI INDIRETTI ---
                cpm_nodes = calcola_cpm(st.session_state.wbs_data)
                is_wp_critical = cpm_nodes.get(wp_id, {}).get('is_critical', False)
                
                min_date = pd.to_datetime(df_reale_calc['Inizio_Previsto']).min().date()
                max_date = pd.to_datetime(df_reale_calc['Fine_Prevista']).max().date()
                
                if pd.notna(min_date) and pd.notna(max_date):
                    giorni_pianificati = (max_date - min_date).days
                    tot_pv = df_reale_calc['PV'].sum()
                    tot_ev = df_reale_calc['EV'].sum()
                    spi_attuale = tot_ev / tot_pv if tot_pv > 0 else 1.0
                    
                    giorni_stimati_attuali = int(giorni_pianificati / spi_attuale) if spi_attuale > 0 else giorni_pianificati
                    giorni_stimati_attuali = min(giorni_stimati_attuali, giorni_pianificati * 3)
                    
                    data_fine_attuale = min_date + pd.Timedelta(days=giorni_stimati_attuali)
                    
                    # Applica i giorni solo se la voce è critica
                    if is_wp_critical:
                        data_fine_simulata = data_fine_attuale + pd.Timedelta(days=var_giorni)
                        giorni_delta = var_giorni
                    else:
                        data_fine_simulata = data_fine_attuale
                        giorni_delta = 0
                else:
                    data_fine_attuale = data_fine_simulata = oggi
                    giorni_delta = 0
                    
                # Calcolo separato per gli Indiretti
                impatto_indiretti = giorni_delta * costo_indiretto_gg
                costo_totale_effettivo = eac_simulato + impatto_indiretti
                
                # --- MESSAGGI DI ALLERTA INTELLIGENTI ---
                if var_giorni != 0 and not is_wp_critical:
                    st.warning(f"⚠️ **Attenzione:** Stai modificando un'attività NON CRITICA. Cambiare i tempi di '{wp_scelto.split(' - ')[1]}' non sposterà la data di fine cantiere, quindi non inciderà sui costi indiretti.")
                elif var_giorni < 0 and is_wp_critical:
                    st.success(f"✅ **Crashing Efficace:** L'attività è critica. Anticipando di {abs(var_giorni)} giorni la consegna, generi un risparmio extra (indiretto) di € {abs(impatto_indiretti):,.0f}.")
                elif var_giorni > 0 and is_wp_critical:
                    st.error(f"🚨 **Allerta Ritardo:** Un ritardo critico di {var_giorni} giorni farà slittare il cantiere, aggiungendo € {impatto_indiretti:,.0f} ai tuoi costi fissi.")
                
                c_res1, c_res2, c_res3 = st.columns([1, 1, 1.5])
                
                with c_res1:
                    st.metric("EAC Simulato (Puro EVM)", f"€ {eac_simulato:,.0f}", delta=f"Deriva CPI: € {eac_simulato - eac_attuale:,.0f}", delta_color="inverse")
                    
                    if giorni_delta < 0:
                        label_delta_giorni = f"Anticipo: {abs(giorni_delta)} gg"
                        colore_delta = "normal"
                    elif giorni_delta > 0:
                        label_delta_giorni = f"Ritardo: {giorni_delta} gg"
                        colore_delta = "inverse"
                    else:
                        label_delta_giorni = "Nessuna variazione globale"
                        colore_delta = "off"
                        
                    st.metric("Nuova Data di Consegna", data_fine_simulata.strftime('%d/%m/%Y') if pd.notna(data_fine_simulata) else "N/D", delta=label_delta_giorni, delta_color=colore_delta)
                
                with c_res2:
                    st.metric("Impatto Costi Indiretti", f"€ {impatto_indiretti:,.0f}", delta="Risparmio da anticipo" if impatto_indiretti < 0 else ("Penale / Spesa extra" if impatto_indiretti > 0 else "Nessun impatto"), delta_color="inverse")
                    st.metric("Costo Finale Effettivo", f"€ {costo_totale_effettivo:,.0f}", delta=f"Delta vs Attuale: € {costo_totale_effettivo - eac_attuale:,.0f}", delta_color="inverse")

                with c_res3:
                    fig_sim = go.Figure()
                    
                    # Traiettoria Attuale (Rossa)
                    fig_sim.add_trace(go.Scatter(
                        x=[oggi, data_fine_attuale], 
                        y=[ac_attuale_tot, eac_attuale], 
                        mode='lines+markers+text', 
                        name='EAC Attuale', 
                        line=dict(color='red', dash='dash', width=2), 
                        text=["", f"€ {eac_attuale:,.0f}"], 
                        textposition="bottom right"
                    ))
                    
                    # Traiettoria Simulata (Blu) - FIX: Ora parte dallo stesso punto di oggi (AC Attuale) e punta dritta al nuovo traguardo!
                    fig_sim.add_trace(go.Scatter(
                        x=[oggi, data_fine_simulata], 
                        y=[ac_attuale_tot, costo_totale_effettivo], 
                        mode='lines+markers+text', 
                        name='EAC + Indiretti', 
                        line=dict(color='blue', dash='solid', width=3), 
                        text=["", f"€ {costo_totale_effettivo:,.0f}"], 
                        textposition="top left"
                    ))
                    
                    fig_sim.update_layout(
                        title="Impatto Strategico Globale", 
                        height=280, 
                        margin=dict(l=10, r=10, t=35, b=10), 
                        yaxis_title="Costo (€)", 
                        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
                    )
                    st.plotly_chart(fig_sim, use_container_width=True)
                    
    # ---------------------------------------------------------
        
        st.subheader("3. Stampa Verbale di Direzione Lavori")
        
        # 1. Filtri per scegliere cosa stampare
        col_f1, col_f2 = st.columns([1, 2])
        filtro_stampa = col_f1.radio("Quali interventi includere nel verbale?", ["Tutti i registrati", "Solo l'ultimo inserito", "Intervallo di date"])
        
        df_stampa = st.session_state.capa_data.copy()
        if not df_stampa.empty:
            df_stampa['Data_Apertura'] = pd.to_datetime(df_stampa['Data_Apertura']).dt.date
            
            if filtro_stampa == "Solo l'ultimo inserito":
                df_stampa = df_stampa.tail(1)
            elif filtro_stampa == "Intervallo di date":
                d_start = col_f2.date_input("Da data:", value=pd.Timestamp.today().date())
                d_end = col_f2.date_input("A data:", value=pd.Timestamp.today().date())
                df_stampa = df_stampa[(df_stampa['Data_Apertura'] >= d_start) & (df_stampa['Data_Apertura'] <= d_end)]

        # 2. Generazione automatica in background del documento WORD
        df_evm_rep = calcola_evm(get_foglie(st.session_state.wbs_data), pd.Timestamp.today().date())
        tot_ev_rep = df_evm_rep['EV'].sum()
        tot_ac_rep = df_evm_rep['AC_Costo_Reale'].sum()
        tot_pv_rep = df_evm_rep['PV'].sum()
        cpi_rep = tot_ev_rep / tot_ac_rep if tot_ac_rep > 0 else 1.0
        spi_rep = tot_ev_rep / tot_pv_rep if tot_pv_rep > 0 else 1.0
        eac_rep = df_evm_rep['EAC'].sum()
        
        doc = Document()
        doc.add_heading('VERBALE DI DIREZIONE LAVORI', 0)
        doc.add_paragraph(f"Progetto: {st.session_state.nome_progetto_attivo}")
        doc.add_paragraph(f"Data emissione verbale: {pd.Timestamp.today().strftime('%d/%m/%Y')}")
        
        doc.add_heading('1. Stato Avanzamento Lavori (EVM)', level=1)
        p = doc.add_paragraph()
        p.add_run(f"CPI (Efficienza Costi): {cpi_rep:.2f}\n").bold = True
        p.add_run(f"SPI (Efficienza Tempi): {spi_rep:.2f}\n").bold = True
        p.add_run(f"Costo Finale Stimato (EAC): {eac_rep:,.2f} Euro").bold = True
        doc.add_paragraph("Nota: Un indicatore inferiore a 1.00 indica un superamento del budget o un ritardo sui tempi.")
        
        doc.add_heading('2. Disposizioni e Azioni (CAPA)', level=1)
        
        if not df_stampa.empty:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Data'
            hdr_cells[1].text = 'Attività WBS'
            hdr_cells[2].text = 'Tipo'
            hdr_cells[3].text = 'Descrizione Intervento'
            hdr_cells[4].text = 'Stato'
            
            for _, row in df_stampa.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Data_Apertura'])
                row_cells[1].text = str(row['ID_WBS_Rif'])
                row_cells[2].text = str(row['Tipo_Azione'])
                row_cells[3].text = str(row['Descrizione']) + f"\n(Assegnato: {row['Responsabile_OBS']})"
                row_cells[4].text = str(row['Stato'])
        else:
            doc.add_paragraph("Nessun intervento registrato nel periodo selezionato.")
            
        doc.add_paragraph("\n\nFirma Direzione Lavori\n_________________________")
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # 3. Tasto nativo di download (fuori da altri bottoni!)
        st.download_button(
            label="⬇️ Scarica il file Word pronto per la firma",
            data=buffer,
            file_name=f"Verbale_{st.session_state.nome_progetto_attivo}_{pd.Timestamp.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
        
    # ---------------------------------------------------
    # --- TAB 8: MATRICE DEI RISCHI (RISK MANAGEMENT) ---
    # ---------------------------------------------------
    
    with tab8:
        st.header("Matrice di Rischio Strategico")
        st.markdown("Valuta e mappa i rischi associati alle singole lavorazioni (WBS) assegnando un punteggio da **1 (Minimo)** a **5 (Massimo)**.")
        
        leaf_wbs_rischi = get_foglie(st.session_state.wbs_data)
        wbs_options_rischi = [f"{row['ID_WBS']} - {row['Attività']}" for _, row in leaf_wbs_rischi.iterrows()]
        
        st.subheader("1. Registro dei Rischi (Risk Register)")
        
        edited_rischi = st.data_editor(
            st.session_state.rischi_data, num_rows="dynamic", use_container_width=True, hide_index=True,
            column_config={
                "ID_WBS_Rif": st.column_config.SelectboxColumn("Attività WBS (Rif.)", options=wbs_options_rischi),
                "Descrizione_Rischio": st.column_config.TextColumn("Descrizione / Nome del Rischio", width="large"),
                "Probabilità (1-5)": st.column_config.NumberColumn("Probabilità (1=Rara, 5=Certa)", min_value=1, max_value=5, step=1),
                "Impatto (1-5)": st.column_config.NumberColumn("Impatto (1=Lieve, 5=Critico)", min_value=1, max_value=5, step=1),
                "Stato": st.column_config.SelectboxColumn("Stato", options=["Attivo ▾", "Monitorato ▾", "Mitigato ▾", "Chiuso ▾"])
            }
        )
        
        st.divider()
        st.warning("⚠️ **Ricordati di cliccare il tasto rosso qui sotto dopo aver modificato o aggiunto i rischi!**")
        if st.button("💾 SALVA REGISTRO RISCHI E AGGIORNA GRAFICO", type="primary", use_container_width=True):
            st.session_state.rischi_data = edited_rischi
            st.success("✅ Rischi mappati con successo!")
            st.rerun()
            
        st.divider()
        st.subheader("2. Mappa di Calore Operativa (Heatmap)")
        
        # IL TUO NUOVO INTERRUTTORE
        mostra_nomi = st.toggle("👁️ Mostra nomi completi delle attività sul grafico", value=False)
        
        df_plot = st.session_state.rischi_data.copy()
        df_plot = df_plot.dropna(subset=['Probabilità (1-5)', 'Impatto (1-5)'])
        
        if not df_plot.empty:
            df_plot['Probabilità (1-5)'] = pd.to_numeric(df_plot['Probabilità (1-5)'])
            df_plot['Impatto (1-5)'] = pd.to_numeric(df_plot['Impatto (1-5)'])
            df_plot['Punteggio'] = df_plot['Probabilità (1-5)'] * df_plot['Impatto (1-5)']
            
            fig_risk = go.Figure()
            
            # --- 4 QUADRANTI COLORATI ---
            fig_risk.add_shape(type="rect", x0=0.5, y0=0.5, x1=3, y1=3, fillcolor="#E8F5E9", line_width=0, layer="below")
            fig_risk.add_shape(type="rect", x0=0.5, y0=3, x1=3, y1=5.5, fillcolor="#FFFDE7", line_width=0, layer="below")
            fig_risk.add_shape(type="rect", x0=3, y0=0.5, x1=5.5, y1=3, fillcolor="#FFFDE7", line_width=0, layer="below")
            fig_risk.add_shape(type="rect", x0=3, y0=3, x1=5.5, y1=5.5, fillcolor="#FFEBEE", line_width=0, layer="below")
            
            # Linee tratteggiate a Croce
            fig_risk.add_hline(y=3, line_dash="dash", line_color="gray", line_width=2)
            fig_risk.add_vline(x=3, line_dash="dash", line_color="gray", line_width=2)
            
            # --- TRACCE FANTASMA PER LA LEGENDA GRAFICA ---
            fig_risk.add_trace(go.Scatter(x=[None], y=[None], mode='markers', marker=dict(size=18, color='#E8F5E9', symbol='square', line=dict(color='gray', width=1)), name='Safe-Zone (Rischio Controllato)'))
            fig_risk.add_trace(go.Scatter(x=[None], y=[None], mode='markers', marker=dict(size=18, color='#FFFDE7', symbol='square', line=dict(color='gray', width=1)), name='Area di Attenzione'))
            fig_risk.add_trace(go.Scatter(x=[None], y=[None], mode='markers', marker=dict(size=18, color='#FFEBEE', symbol='square', line=dict(color='gray', width=1)), name='Area Critica'))
            
            # Logica dell'interruttore: Testo Lungo vs ID Corto
            if mostra_nomi:
                etichette_punti = df_plot['ID_WBS_Rif'].apply(lambda x: str(x) if pd.notna(x) else "")
            else:
                etichette_punti = df_plot['ID_WBS_Rif'].apply(lambda x: str(x).split(' - ')[0] if pd.notna(x) else "")

            # --- I PUNTI DEI RISCHI (COLORATI IN BASE ALLO STATO) ---
            colori_stato = {"Attivo ▾": "#F44336", "Monitorato ▾": "#FF9800", "Mitigato ▾": "#4CAF50", "Chiuso ▾": "#9E9E9E"}
            df_plot['Colore_Punto'] = df_plot['Stato'].map(colori_stato).fillna("#607D8B")
            
            fig_risk.add_trace(go.Scatter(
                x=df_plot['Probabilità (1-5)'], y=df_plot['Impatto (1-5)'],
                mode='markers+text',
                text=etichette_punti,
                textposition="top center",
                textfont=dict(size=12, color='DarkSlateGrey', family="Arial Black"),
                marker=dict(size=18, color=df_plot['Colore_Punto'], line=dict(width=2, color='white')),
                showlegend=False,
                hovertemplate="<b>WBS: %{text}</b><br>Rischio: %{customdata[0]}<br>Probabilità: %{x}<br>Impatto: %{y}<br>Punteggio: %{customdata[2]}<br>Stato: %{customdata[1]}<extra></extra>",
                customdata=df_plot[['Descrizione_Rischio', 'Stato', 'Punteggio']]
            ))
            
            # --- IL PALLINO ROSSO GIGANTE DEL RISCHIO MEDIO ---
            media_prob = df_plot['Probabilità (1-5)'].mean()
            media_imp = df_plot['Impatto (1-5)'].mean()
            
            fig_risk.add_trace(go.Scatter(
                x=[media_prob], 
                y=[media_imp],
                mode='markers+text',
                text=["RISCHIO MEDIO"],
                textposition="bottom center",
                textfont=dict(size=14, color='DarkRed', family="Arial Black"),
                marker=dict(size=35, color='rgba(255, 0, 0, 0.7)', line=dict(width=3, color='DarkRed')),
                showlegend=False,
                hovertemplate="<b>RISCHIO GLOBALE MEDIO</b><br>Probabilità Media: %{x:.2f}<br>Impatto Medio: %{y:.2f}<extra></extra>"
            ))
            
            fig_risk.update_layout(
                xaxis=dict(title="<b>Probabilità di Accadimento (1-5)</b>", range=[0.5, 5.5], dtick=1, gridcolor='white', zeroline=False),
                yaxis=dict(title="<b>Impatto sul Progetto (1-5)</b>", range=[0.5, 5.5], dtick=1, gridcolor='white', zeroline=False),
                height=650, plot_bgcolor='white', margin=dict(l=60, r=40, t=40, b=60),
                legend=dict(yanchor="top", y=0.99, xanchor="left", x=0.01, bgcolor="rgba(255, 255, 255, 0.9)", bordercolor="lightgray", borderwidth=1)
            )
            
            st.plotly_chart(fig_risk, use_container_width=True)
            
            st.markdown("**Legenda Stato Rischi:** 🔴 `Attivo` | 🟠 `Monitorato` | 🟢 `Mitigato` | ⚪ `Chiuso`")
        else:
            st.info("ℹ️ Compila i valori numerici (da 1 a 5) nella colonna Probabilità e Impatto della tabella qui sopra per generare la matrice.")

    # ---------------------------------------
    # --- TAB 9: COMUNICAZIONI E VARIANTI ---
    # ---------------------------------------
    with tab9: 
        st.header("📩 Registro Varianti e Comunicazioni")
        st.markdown("*Ogni richiesta viene storicizzata con marcatura temporale. Le richieste inviate non possono essere modificate dall'autore, garantendo un Audit Trail.*")
        
        # --- SEZIONE 1: IL MITTENTE APRE IL TICKET ---
        with st.expander("➕ Apri un nuovo Ticket (Riservato a DL / CSE / Impresa)", expanded=False):
            with st.form("form_nuovo_ticket"):
                c1, c2 = st.columns(2)
                
                # Creiamo la lista delle WBS dal tab 1 per il menu a tendina
                lista_wbs = []
                if 'wbs_data' in st.session_state and not st.session_state.wbs_data.empty:
                    lista_wbs = [f"{row['ID_WBS']} - {row['Attività']}" for _, row in st.session_state.wbs_data.iterrows() if pd.notna(row['ID_WBS'])]
                
                wbs_rif = c1.selectbox("WBS di Riferimento", options=["Nessuna"] + lista_wbs)
                autore = c2.selectbox("Ruolo Autore", options=["Direttore dei Lavori (DL)", "Coordinatore Sicurezza (CSE)", "Impresa Appaltatrice"])
                
                tipo = c1.selectbox("Tipologia di Richiesta", options=["Richiesta di Variante", "Richiesta di Chiarimento Progettuale", "Segnalazione Imprevisto/Ostacolo", "Ordine di Servizio"])
                descrizione = st.text_area("Descrizione dettagliata (Attenzione: Non modificabile dopo l'invio)")
                
                submit_ticket = st.form_submit_button("Invia Ticket al RUP", type="primary")
                
                if submit_ticket:
                    if wbs_rif == "Nessuna" or not descrizione.strip():
                        st.warning("⚠️ Seleziona una WBS e inserisci una descrizione prima di inviare.")
                    else:
                        import datetime
                        # Creazione ID progressivo (es. TCK-001)
                        nuovo_id = f"TCK-{len(st.session_state.tickets_data)+1:03d}"
                        wbs_clean = wbs_rif.split(' - ')[0] # Estrae solo il codice numerico
                        
                        nuovo_ticket = {
                            'ID_Ticket': nuovo_id,
                            'ID_WBS_Rif': wbs_clean,
                            'Autore': autore,
                            'Data_Apertura': datetime.datetime.now().strftime("%d/%m/%Y %H:%M"),
                            'Tipologia': tipo,
                            'Descrizione': descrizione,
                            'Stato': 'In attesa ⏳',
                            'Risposta_RUP': '',
                            'Data_Chiusura': None,
                            'Variazione_Costi': None,
                            'Variazione_Tempi': None,
                            'Variante_Applicata': False
                        }
                        
                        st.session_state.tickets_data = pd.concat([st.session_state.tickets_data, pd.DataFrame([nuovo_ticket])], ignore_index=True)
                        st.success("✅ Ticket registrato nell'Audit Trail con successo!")
                        st.rerun()
        
        st.divider()
        
        # --- SEZIONE 2: IL RUP GESTISCE E RISPONDE ---
        st.subheader("📋 Audit Trail: Cruscotto di Risposta RUP")
        
        if not st.session_state.tickets_data.empty:
            df_tickets = st.session_state.tickets_data.copy()
            
            # BLOCCAGGIO INCORRUTTIBILE: Variante_Applicata viene nascosta all'utente
            colonne_bloccate = ['ID_Ticket', 'ID_WBS_Rif', 'Autore', 'Data_Apertura', 'Tipologia', 'Descrizione', 'Data_Chiusura', 'Variante_Applicata']
            
            edited_tickets = st.data_editor(
                df_tickets,
                disabled=colonne_bloccate,
                num_rows="dynamic", 
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Stato": st.column_config.SelectboxColumn("Stato Richiesta", options=["In attesa ⏳", "Approvato ✅", "Respinto ❌"], required=True),
                    "Risposta_RUP": st.column_config.TextColumn("Prescrizioni / Risposta RUP"),
                    "Variazione_Costi": st.column_config.NumberColumn("Variazione Costi (€)", format="€ %.2f", help="Aggiunge o sottrae budget alla WBS"),
                    "Variazione_Tempi": st.column_config.NumberColumn("Variazione Tempi (gg)", help="Sposta la Fine Prevista della WBS"),
                    "Variante_Applicata": None # Nasconde la colonna di sistema
                }
            )
            
            if st.button("💾 Registra Risposte RUP"):
                import datetime
                
                # 1. Prima di tutto, salviamo a forza i testi modificati nel database dei ticket
                for idx, row in edited_tickets.iterrows():
                    st.session_state.tickets_data.at[idx, 'Stato'] = row.get('Stato', 'In attesa ⏳')
                    st.session_state.tickets_data.at[idx, 'Risposta_RUP'] = row.get('Risposta_RUP', '')
                    st.session_state.tickets_data.at[idx, 'Variazione_Costi'] = row.get('Variazione_Costi')
                    st.session_state.tickets_data.at[idx, 'Variazione_Tempi'] = row.get('Variazione_Tempi')
                    
                    if row.get('Stato') != 'In attesa ⏳' and pd.isna(st.session_state.tickets_data.at[idx, 'Data_Chiusura']):
                        st.session_state.tickets_data.at[idx, 'Data_Chiusura'] = datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
                    elif row.get('Stato') == 'In attesa ⏳':
                        st.session_state.tickets_data.at[idx, 'Data_Chiusura'] = None

                # 2. Ora applichiamo le varianti economiche e temporali direttamente sul database WBS
                for idx, row in edited_tickets.iterrows():
                    t_tipo = st.session_state.tickets_data.at[idx, 'Tipologia']
                    t_stato = st.session_state.tickets_data.at[idx, 'Stato']
                    
                    if t_tipo == 'Richiesta di Variante' and t_stato == 'Approvato ✅':
                        c_val = pd.to_numeric(st.session_state.tickets_data.at[idx, 'Variazione_Costi'], errors='coerce')
                        t_val = pd.to_numeric(st.session_state.tickets_data.at[idx, 'Variazione_Tempi'], errors='coerce')
                        
                        v_app = st.session_state.tickets_data.at[idx, 'Variante_Applicata']
                        applicata = v_app if isinstance(v_app, bool) else str(v_app).strip().lower() in ['true', '1', 't', 'y', 'yes']
                        
                        if (pd.notna(c_val) or pd.notna(t_val)) and not applicata:
                            c_clean = float(c_val) if pd.notna(c_val) else 0.0
                            t_clean = int(t_val) if pd.notna(t_val) else 0
                            
                            # FIX CRITICO: Estrazione purificata dell'ID WBS (rimuove testi e spazi accidentali)
                            wbs_target_raw = str(st.session_state.tickets_data.at[idx, 'ID_WBS_Rif']).strip()
                            wbs_target = wbs_target_raw.split(' - ')[0].strip()
                            
                            # Cerchiamo la riga esatta nell'albero WBS
                            mask = st.session_state.wbs_data['ID_WBS'].astype(str).str.strip() == wbs_target
                            
                            if mask.any():
                                i_w = st.session_state.wbs_data.index[mask][0]
                                
                                # A) Incrementa il Budget (BAC)
                                if c_clean != 0:
                                    b_att = pd.to_numeric(st.session_state.wbs_data.at[i_w, 'BAC_Budget'], errors='coerce')
                                    st.session_state.wbs_data.at[i_w, 'BAC_Budget'] = (b_att if pd.notna(b_att) else 0.0) + c_clean
                                
                                # B) Sposta in avanti la Data Fine Prevista
                                if t_clean != 0:
                                    f_att = pd.to_datetime(st.session_state.wbs_data.at[i_w, 'Fine_Prevista'], errors='coerce')
                                    if pd.notna(f_att):
                                        st.session_state.wbs_data.at[i_w, 'Fine_Prevista'] = (f_att + pd.Timedelta(days=t_clean)).date()
                                        
                                    f_eff = st.session_state.wbs_data.at[i_w, 'Fine_Effettiva']
                                    if pd.notna(f_eff) and str(f_eff).strip() not in ['', 'NaT', 'None', 'nan']:
                                        st.session_state.wbs_data.at[i_w, 'Fine_Effettiva'] = (pd.to_datetime(f_eff) + pd.Timedelta(days=t_clean)).date()

                                # C) Mette il sigillo per evitare doppi conteggi
                                st.session_state.tickets_data.at[idx, 'Variante_Applicata'] = True
                                st.toast(f"✅ Variante iniettata: WBS {wbs_target} aggiornata (+{c_clean}€, +{t_clean}gg)!", icon="⚙️")
                            else:
                                st.error(f"❌ ERRORE: La WBS '{wbs_target}' non esiste nell'albero! Variante annullata.")
                                
                # 3. Ricalcolo globale della gerarchia, dei padri e dei motori EVM/CPM
                st.session_state.wbs_data = aggiorna_gerarchia(st.session_state.wbs_data)
                
                # Pulizia della cache dei widget per forzare il refresh visivo del Tab 1
                for k in list(st.session_state.keys()):
                    if k.startswith("editor_wbs_"):
                        del st.session_state[k]
                        
                st.success("✅ Registro aggiornato! Il cronoprogramma e l'EVM sono stati ricalcolati.")
                import time
                time.sleep(2)
                st.rerun()
        else:
            st.info("Nessuna comunicazione o variante registrata.")

        # ========================================================
        # --- SEZIONE 3: ESPORTAZIONE E STAMPA AUDIT TRAIL ---
        # ========================================================
        st.divider()
        st.subheader("🖨️ 3. Stampa Registro Audit Trail")
        
        # 1. Filtri per scegliere cosa stampare
        col_fa1, col_fa2 = st.columns([1, 2])
        filtro_stampa_audit = col_fa1.radio("Quali Ticket includere nel report?", ["Tutti i registrati", "Solo Varianti Approvate ✅", "Intervallo di date"])
        
        df_stampa_audit = st.session_state.tickets_data.copy()
        if not df_stampa_audit.empty:
            # Creiamo una colonna data fittizia per filtrare (Data_Apertura è in formato testuale GG/MM/YYYY HH:MM)
            df_stampa_audit['Data_Filtro'] = pd.to_datetime(df_stampa_audit['Data_Apertura'], format="%d/%m/%Y %H:%M", errors='coerce').dt.date
            
            if filtro_stampa_audit == "Solo Varianti Approvate ✅":
                df_stampa_audit = df_stampa_audit[(df_stampa_audit['Stato'] == 'Approvato ✅') & (df_stampa_audit['Tipologia'] == 'Richiesta di Variante')]
            elif filtro_stampa_audit == "Intervallo di date":
                da_data = col_fa2.date_input("Da data:", value=pd.Timestamp.today().date(), key="data_da_audit")
                a_data = col_fa2.date_input("A data:", value=pd.Timestamp.today().date(), key="data_a_audit")
                df_stampa_audit = df_stampa_audit[(df_stampa_audit['Data_Filtro'] >= da_data) & (df_stampa_audit['Data_Filtro'] <= a_data)]

        # 2. Generazione automatica del documento Word
        if not df_stampa_audit.empty:
            doc_audit = Document()
            doc_audit.add_heading('REGISTRO AUDIT TRAIL - COMUNICAZIONI E VARIANTI', 0)
            doc_audit.add_paragraph(f"Progetto: {st.session_state.nome_progetto_attivo}")
            doc_audit.add_paragraph(f"Data emissione registro: {pd.Timestamp.today().strftime('%d/%m/%Y')}")
            
            for _, row in df_stampa_audit.iterrows():
                doc_audit.add_heading(f"Ticket {row['ID_Ticket']} - WBS {row['ID_WBS_Rif']} ({row['Tipologia']})", level=2)
                
                p = doc_audit.add_paragraph()
                p.add_run(f"Autore: {row['Autore']} - Aperto il: {row['Data_Apertura']}\n").bold = True
                p.add_run(f"Stato Attuale: {row['Stato']}").bold = True
                
                if pd.notna(row.get('Data_Chiusura')) and str(row.get('Data_Chiusura')).strip() not in ['None', 'nan', '']:
                    p.add_run(f" (in data {row['Data_Chiusura']})")
                
                doc_audit.add_paragraph(f"Descrizione Richiesta:\n{row['Descrizione']}")
                
                if str(row.get('Risposta_RUP', '')).strip() not in ['None', 'nan', '']:
                    doc_audit.add_paragraph(f"Risposta / Prescrizioni RUP:\n{row['Risposta_RUP']}")
                    
                # Se ci sono variazioni economiche/temporali valorizzate
                if pd.notna(row.get('Variazione_Costi')) or pd.notna(row.get('Variazione_Tempi')):
                    p_var = doc_audit.add_paragraph()
                    p_var.add_run("Impatto Finanziario/Temporale Approvato:\n").bold = True
                    if pd.notna(row.get('Variazione_Costi')):
                        p_var.add_run(f"- Costo Extra: € {row['Variazione_Costi']:,.2f}\n")
                    if pd.notna(row.get('Variazione_Tempi')):
                        p_var.add_run(f"- Slittamento Tempi: {row['Variazione_Tempi']} giorni")
                
                doc_audit.add_paragraph("_" * 60) # Linea di separazione visiva
                
            buffer_audit = BytesIO()
            doc_audit.save(buffer_audit)
            buffer_audit.seek(0)
            
            st.download_button(
                label="⬇️ Scarica Registro Audit Trail",
                data=buffer_audit,
                file_name=f"Audit_Trail_{st.session_state.nome_progetto_attivo}_{pd.Timestamp.today().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )
        else:
            if not st.session_state.tickets_data.empty:
                st.warning("Nessun Ticket corrisponde ai filtri selezionati.")
    
    # ========================================================
    # --- TAB 10: MANUALE OPERATIVO, FORMULARIO & FAQ V.2.0 ---
    # ========================================================
    with tab10:
        st.header("📖 Manuale Operativo & Knowledge Base (Versione 2.0)")
        st.markdown("Benvenuto nella centrale di controllo della documentazione di cantiere. Questo manuale guida l'utente attraverso l'architettura dei dati, le formule matematiche e le logiche di automazione integrate nell'applicazione.")

        # Sotto-sezioni del manuale per massima pulizia visiva
        t_sec1, t_sec2, t_sec3, t_sec4, t_sec5 = st.tabs([
            "🧭 Flusso di Lavoro", 
            "📚 Glossario Tecnico",
            "📐 Formulario EVM & Finanza", 
            "⚠️ FAQ & Falsi Ingippi", 
            "🛠️ Roadmap Versione 2.0"
        ])

        # --- SEZIONE 1: FLUSSO DI LAVORO ---
        with t_sec1:
            st.subheader("Il Ciclo di Vita del Progetto nell'App")
            st.markdown("Il software è progettato in modo che le informazioni viaggino automaticamente tra le varie sezioni, creando un ciclo continuo di pianificazione, misurazione, allerta e correzione.")
            
            st.markdown("""
        ### 🕸️ Architettura del Sistema (Flussi di Dati)
        """)
            
            st.markdown("""
        ```mermaid
        graph TD
            %% Definizione Stili Personalizzati
            classDef planning fill:#E3F2FD,stroke:#1E88E5,stroke-width:2px,color:black;
            classDef control fill:#E8F5E9,stroke:#43A047,stroke-width:2px,color:black;
            classDef alerts fill:#FFEBEE,stroke:#E53935,stroke-width:2px,color:black;
            classDef finance fill:#FFF3E0,stroke:#FB8C00,stroke-width:2px,color:black;
            classDef matrix fill:#E0F7FA,stroke:#00BCD4,stroke-width:2px,color:black;

            subgraph FASE_1 [FASE 1: PIANIFICAZIONE E RISORSE]
                T2[Tab 2: OBS & Risorse]:::planning
                T1[Tab 1: WBS & Budget]:::planning
                T3[Tab 3: Matrice RACI / Incroci]:::matrix
                T4[Tab 4: Gantt & Scadenzario]:::planning
                
                T2 -- "Fornisce Soggetti" --> T3
                T1 -- "Fornisce Attività" --> T3
                T1 -- "Budget e Avanzamento" --> T5
                T3 -. "Valida Assegnazioni" .-> T1
                T1 -- "Date, Predecessori" --> T4
                T8 -- "Alert sul percorco critico" --> T3
                T1 -- "Genera Allerta" --> Scad[Scadenzario Vincoli]:::alerts
                Scad -. "Visibile in" .-> T4
            end

            subgraph FASE_2 [FASE 2: GESTIONE IMPREVISTI E CORREZIONI]
                T8[Tab 8: Rischi]:::finance
                T7[Tab 7: CAPA]:::alerts
                T9[Tab 9: Varianti RUP]:::alerts
                
                T9 == "Inietta Soldi e Giorni" ==> T1
                T8 -- "Riserva Monetaria" --> T5
                T5 -. "Indice SPI colora le barre" .-> T4
                T4 -- "Calcola Percorso Critico" --> T8
                T8 -. "Minaccia Scadenze Critiche" .-> T4
                T8 -. "Evidenzia Criticità Risorse" .-> T3
            end

            subgraph FASE_3 [FASE 3: MOTORI FINANZIARI E CONTROLLO]
                T6[Tab 6: Registro Contabile]:::finance
                T5[Tab 5: Motore EVM & Cash Flow]:::control
                
                
                T7 -- "Genera Costi Tossici" --> T6
                T7 -. "Blocca WBS al 99%" .-> T1
                T6 -- "Genera" --> Flussi[Andamento Flussi di Cassa]:::alerts
                T6 -- "Costi Reali e SAL" --> T5
            end
            
            %% Stile per il riquadro della Fase 3 (Lilla tenue)
            style FASE_3 fill:#F3E5F5,stroke:#8E24AA,stroke-width:2px,stroke-dasharray: 4 4
        ```
        """)
            
           # --- LEGENDA STREAMLIT SOTTO IL GRAFICO ---
            st.markdown("<br>", unsafe_allow_html=True) # Spazio vuoto
            st.markdown("#### 🎨 Legenda Nodi e Connessioni")
            
            legenda_html = """
            <!-- Legenda Nodi (Box Colorati) -->
            <div style="display: flex; flex-wrap: wrap; gap: 10px; justify-content: center; margin-bottom: 15px;">
                <div style="background-color: #E3F2FD; border: 2px solid #1E88E5; padding: 5px 15px; border-radius: 8px; color: black; font-weight: bold;">
                    📘 Pianificazione (Input)
                </div>
                <div style="background-color: #E8F5E9; border: 2px solid #43A047; padding: 5px 15px; border-radius: 8px; color: black; font-weight: bold;">
                    📗 Controllo (EVM)
                </div>
                <div style="background-color: #FFF3E0; border: 2px solid #FB8C00; padding: 5px 15px; border-radius: 8px; color: black; font-weight: bold;">
                    📙 Finanza & Rischi
                </div>
                <div style="background-color: #FFEBEE; border: 2px solid #E53935; padding: 5px 15px; border-radius: 8px; color: black; font-weight: bold;">
                    📕 Allarmi & Blocchi
                </div>
                <div style="background-color: #E0F7FA; border: 2px solid #00BCD4; padding: 5px 15px; border-radius: 8px; color: black; font-weight: bold;">
                    📘 Matrici (RACI)
                </div>
            </div>
            
            <!-- Legenda Frecce (Tipologia di Connessione) -->
            <div style="display: flex; flex-wrap: wrap; gap: 20px; justify-content: center; margin-bottom: 20px; font-size: 0.95em;">
                <div style="display: flex; align-items: center; gap: 8px;">
                    <span style="font-weight: 900; font-family: monospace; font-size: 1.2em;">──►</span> 
                    <span><b>Flusso Dati Standard</b> (Lettura o invio normale di dati)</span>
                </div>
                <div style="display: flex; align-items: center; gap: 8px;">
                    <span style="font-weight: 900; font-family: monospace; font-size: 1.2em; letter-spacing: 2px;">- -►</span> 
                    <span><b>Automazione / Allerta</b> (Controlli logici, allerte o blocchi del sistema)</span>
                </div>
                <div style="display: flex; align-items: center; gap: 8px;">
                    <span style="font-weight: 900; font-family: monospace; font-size: 1.2em;">══►</span> 
                    <span><b>Iniezione Strutturale</b> (Modifica diretta e irreversibile delle fondamenta del progetto)</span>
                </div>
            </div>
            """
            st.markdown(legenda_html, unsafe_allow_html=True)
        
            with st.expander("📝 Flusso di Lavoro (Input Dati)"):
                st.markdown("""
                 1. **Tab 2 (OBS & Risorse):** Inserisci le imprese, le maestranze e le attrezzature disponibili. Sono i soggetti che animeranno il cantiere.
                 2. **Tab 1 (WBS - Lavorazioni):** Struttura la gerarchia delle attività. Assegna i budget (BAC), le date, i predecessori, la risorsa responsabile e imposta gli eventuali **Vincoli Burocratici** (es. Genio Civile).
                 3. **Tab 8 (Gestione Rischi):** Mappa i rischi di cantiere, assegnando probabilità e impatto per calcolare matematicamente il fondo imprevisti.
                 4. **Tab 4 (Gantt & Scadenzario):** Controlla l'allineamento temporale (Gantt EVM) e monitora le scadenze amministrative per non bloccare il cantiere.
                 5. **Tab 6 (Gestione Finanziaria):** 
                    * *Uscite:* Registra fatture e costi reali associandoli alle WBS.
                    * *Entrate:* Emetti e traccia i SAL certificati e pagati dalla committenza.
                 6. **Tab 7 (Qualità & CAPA):** Apri azioni correttive (Non-Conformità) qualora qualcosa non rispetti gli standard, bloccando l'avanzamento dei lavori fallati.
                 7. **Tab 9 (Comunicazioni RUP & Varianti):** Gestisci l'approvazione formale di costi extra e proroghe temporali tramite ticket sigillati (Audit Trail).
                 8. **Tab 5 & Radar (GIANFRY CONSIGLIA):** Monitora il cruscotto di controllo per verificare le proiezioni a finire (EVM), l'esposizione di cassa (Cash Flow) e le allerte di sovraccarico.
                """)
            
            with st.expander(" 📌 Legenda dei Flussi Automatici (Il Motore del Software)"):
                st.markdown("""
                 * **Da Tab 1 a Tab 4 (Scadenzario Amministrativo):** Impostare un "Vincolo Burocratico" nel Tab 1 genera automaticamente un alert semaforico nel Tab 4, calcolando i giorni mancanti all'inizio lavori. Spuntare "Assolto" lo archivia in verde.
                 * **Da Tab 9 a Tab 1 e 5 (Motore Varianti):** L'approvazione di una variante nel Tab 9 inietta matematicamente il nuovo Budget (BAC) e i giorni di proroga direttamente nel Tab 1, aggiornando a cascata l'intero albero WBS e ricalibrando l'EVM (SPI/CPI) nel Tab 5. Il ticket viene poi "sigillato" contro le frodi.
                 * **Da Tab 9, le variazioni vanno tutte registrate. Non cancellare le registrazioni, se ci sono modifiche, fosse anche sulla stessa attività, andrà fatta una nuova variazione in coda alle precednenti. Questo per garantire la tracciabilità di ogni azione protocollata.
                 * **Da Tab 8 a Tab 5 (Scudo Finanziario):** I rischi attivi calcolano il Valore Monetario Atteso (EMV), che si somma automaticamente alla stima a finire (EAC) nel Tab 5 per creare la *Contingency Reserve*.
                 * **Da Tab 7 a Tab 1 e 6 (Costi di Non-Qualità):** Un'azione correttiva aperta blocca la WBS al 99% nel Tab 1. Quando chiusa, se genera un costo extra, questo viene contabilizzato automaticamente nel Tab 6 come *Spesa Tossica*.
                 * **Da Tab 6 a Tab 5 (Cash Flow):** Le uscite (costi reali) e le entrate (SAL pagati) alimentano la curva cumulativa e l'indicatore di esposizione finanziaria netta.
                 * **Da Tab 1, 2 a Radar (Sovraccarico):** Il sistema controlla in tempo reale se la stessa risorsa è impegnata su più fronti nello stesso periodo, segnalando l'eventuale conflitto (con opzione di deroga).
                """)

            with st.expander(" 🎯 Guida Operativa agli Imprevisti: Rischi, CAPA o Varianti?"):
                st.markdown("""
                Una delle forze di questo gestionale è la netta separazione tra **previsione**, **errore** e **modifica contrattuale**. In cantiere succedono mille imprevisti: ecco la bussola definitiva per sapere esattamente in quale Tab registrarli per mantenere i conti in perfetto equilibrio.
                
                ### 📊 Matrice Decisionale degli Imprevisti
                
                | Tipologia Evento | Tab | Quando si usa? (La Natura) | Esempio Pratico in Cantiere | Effetto sul Conto (Motore EVM) | Effetto a Cascata nel Sistema |
                | :--- | :---: | :--- | :--- | :--- | :--- |
                | 🔮 **RISCHIO** (Gestione Incertezza) | **8** | Evento **futuro incerto** (potrebbe accadere, ma non è detto). | *Rischio di trovare roccia dura durante le operazioni di scavo.* | Crea la **Riserva (EMV)** | Alza l'**EAC** (Stima a finire) nel cruscotto Tab 5 creando uno "Scudo Finanziario" preventivo, senza toccare i soldi attuali. |
                | 🚨 **CAPA** (Non-Conformità / Errori) | **7** | Errore, difetto o infortunio **già avvenuto** (da sanare o correggere). | *Getto di calcestruzzo errato, armatura da demolire e rifare.* | Crea **Spesa Tossica (AC)** | Blocca la WBS al **99%** (Tab 1). I costi extra di ripristino alzano l'**AC** (Tab 6), bruciando cassa e abbassando l'indice di efficienza (CPI). |
                | 📜 **VARIANTE** (Comunicazione RUP) | **9** | Modifica **formale e approvata** al progetto o al contratto originale. | *La Committenza richiede e approva la costruzione di una stanza extra.* | Cambia il **Budget (BAC)** | Inietta soldi e/o giorni extra nel **Tab 1**. Il "Sigillo" previene manomissioni (Audit Trail). Ricalibra l'intero piano dei tempi e dei costi. |
                
                ---
                
                ### 🛠️ Spiegazione Dettagliata
                
                *   **TAB 8 - GESTIONE RISCHI (La Sfera di Cristallo):** Qui si lavora di prevenzione. Inserisci tutto ciò di cui hai paura *prima* che accada. Il software calcola le probabilità e ti dice quanti soldi tenere da parte. Se il rischio si avvera davvero, lo chiudi qui e ne affronti le conseguenze (aprendo una Variante se il cliente paga, o registrando una spesa se tocca all'impresa).
                *   **TAB 7 - AZIONI CORRETTIVE (Il Pronto Soccorso):** Usalo quando qualcuno ha sbagliato e devi mettere una "toppa". È lo strumento di Qualità e Sicurezza dell'app. Matematicamente è punitivo: ti fa spendere soldi non previsti per riparare un danno, peggiorando le tue statistiche di cantiere.
                *   **TAB 9 - VARIANTI RUP (Il Notaio):** È lo strumento amministrativo per eccellenza. Usalo SOLO quando hai in mano un "Approvato" ufficiale per variare i lavori. Usare questo Tab altera le "fondamenta" stesse del progetto (la Baseline), per questo il sistema lo protegge registrando esattamente chi ha aggiunto i soldi e impedendo di modificarlo di nascosto in un secondo momento.
                """)

        # --- 2. GLOSSARIO EVM ---
        with t_sec2:
            st.subheader("2. Glossario EVM (Earned Value Management)")
        
            c_glos1, c_glos2 = st.columns(2)
            with c_glos1:
                with st.expander("Valori Base (I Pilastri)"):
                    st.markdown("""
                    * **BAC (Budget at Completion):** Il Budget totale pianificato per l'intero progetto o lavorazione.
                    * **PV (Planned Value):** Il valore del lavoro che *dovrebbe* essere stato completato fino ad oggi secondo il cronoprogramma. (Si calcola proiettando il BAC nel tempo).
                    * **EV (Earned Value):** Il valore del lavoro *effettivamente* completato fino ad oggi. È la metrica più importante: indica i soldi che il cantiere ha realmente "guadagnato" producendo.
                    * **AC (Actual Cost):** I costi reali effettivamente sostenuti per il lavoro svolto fino ad oggi (fatture, ore manodopera, materiali).
                    """)
            with c_glos2:
                with st.expander("Indicatori di Performance (KPI)"):
                    st.markdown("""
                    * **CV (Cost Variance):** Varianza dei costi. Se è negativa, stai spendendo più del previsto.
                    * **SV (Schedule Variance):** Varianza dei tempi. Se è negativa, sei in ritardo sul cronoprogramma.
                    * **CPI (Cost Performance Index):** Efficienza dei costi. Valore ideale: ≥ 1.0. Se vale 0.8, significa che per ogni Euro speso stai producendo solo 80 centesimi di valore.
                    * **SPI (Schedule Performance Index):** Efficienza dei tempi. Valore ideale: ≥ 1.0. Se vale 0.9, stai viaggiando al 90% della velocità prevista.
                    """)
                
            with st.expander("Previsioni e Rischio (Proiezioni)"):
                st.markdown("""
                * **EAC (Estimate At Completion):** Costo totale stimato a fine progetto, ricalcolato in base all'efficienza attuale (CPI). Ti dice quanto ti costerà davvero il cantiere se continui a lavorare come stai facendo oggi.
                * **ETC (Estimate To Complete):** I fondi residui necessari per finire il lavoro da oggi in poi.
                * **VAC (Variance At Completion):** Scostamento finale previsto (BAC - EAC). Se è negativo, il progetto si chiuderà in perdita rispetto al budget iniziale.
                * **EMV (Expected Monetary Value):** Valore Monetario Atteso. Trasforma i punteggi di rischio in valuta, creando un *Fondo Imprevisti* dinamico.
                * **EAC Risk-Adjusted:** L'EAC classico sommato al Fondo Imprevisti. È lo scenario finanziario più prudente.
                """)
            
        # --- SEZIONE 3: FORMULARIO ---
        with t_sec3:
            st.subheader("Matematica e Indicatori di Performance (EVM)")
            st.markdown("Il motore calcola in tempo reale lo stato di salute del cantiere utilizzando le metriche standard internazionali dell'**Earned Value Management** e della finanza di progetto:")
            with st.expander("Valori Base (I Pilastri)"):
                st.latex(r"EV = BAC \times \% \text{ Avanzamento Fisico}")
                st.caption("**EV (Earned Value)")
                st.markdown("Planned Value: *sono i costi definiti da progetto/computo, compilati manualmente in WBS*")
                st.markdown("BAC (Budget at Completion): *Il Budget totale pianificato*")
                
                
            with st.expander("Indicatori di Performance (KPI)"):
                st.latex(r"CV = EV - AC")
                st.caption("**CV (Cost Variance):** Variazione dei costi (Valore Guadagnato meno Costo Reale).")
                st.latex(r"SV = EV - PV")
                st.caption("**SV (Schedule Variance)")
                st.latex(r"CPI = \frac{EV}{AC}")
                st.caption("**CPI (Cost Performance Index):** Efficienza economica. Se $< 0.95$, si sta spendendo più del budget previsto.")
                st.latex(r"SPI = \frac{EV}{PV}")
                st.caption("**SPI (Schedule Performance Index):** Efficienza temporale. Se $< 0.95$, il cantiere è in ritardo rispetto al cronoprogramma.")
                
            with st.expander("Previsioni e Rischio (Proiezioni)"):
                st.latex(r"EAC = \frac{BAC}{CPI}")
                st.caption("**EAC (Estimate At Completion)")
                st.latex(r"ETC = EAC - AC")
                st.caption("**ETC (Estimate To Complete)")
                st.latex(r"VAC = BAC - EAC")
                st.caption("**VAC (Variance At Completion)")    
                st.latex(r"\text{EAC Risk-Adjusted} = EAC + \sum (\text{Budget}_\text{WBS} \times \text{Probabilità}_\text{Rischio} \times \text{Impatto}_\text{Rischio})")
                st.caption("**Quanta liquidità manca al competamento in considerazione dei rischi nella loro probabilità e impatto")
                st.latex(r"CF_{netto} = \sum Entrate_{SAL} - \sum Uscite_{AC}")
                st.caption("**Cash Flow Netto:** Esposizione di cassa al netto dei pagamenti ricevuti.")

        # --- SEZIONE 4: CASISTICHE E FALSI INGHIPPI ---
        with t_sec4:
            st.subheader("🔍 Guida pratica ai 'Falsi Ingippi' e Blocchi di Sicurezza")
            st.markdown("In questa sezione spieghiamo i comportamenti automatizzati del software che potrebbero sembrare anomalie, ma che in realtà sono **controlli di sicurezza attivi**.")

            with st.expander("🚧 1. Il mistero del '99%' (Blocco Qualità CAPA)"):
                st.markdown("""
                * **La Situazione:** Hai impostato una lavorazione al `100%` nel Tab 1, ma il sistema la corregge d'ufficio al `99%` e ti mostra un errore rosso.
                * **Perché accade:** C'è una **CAPA (Non-Conformità) aperta o in lavorazione** nel Tab 7 associata a quella specifica WBS. 
                * **La Logica:** Il software impedisce al Direttore Lavori di chiudere contabilmente un'attività finché il problema di qualità o sicurezza non è stato formalmente risolto.
                * **Come sbloccarlo:** Vai nel Tab 7, verifica l'azione correttiva e imposta lo stato su **'Chiuso'**. Torna nel Tab 1 e potrai finalmente certificare il 100%.
                """)

            with st.expander("📉 2. Il grafico del Cash Flow non si aggiorna o mostra solo un punto"):
                st.markdown("""
                * **La Situazione:** Hai inserito i SAL nel Tab 6, ma il grafico a gradoni non mostra le linee o la curva delle uscite è piatta.
                * **Perché accade:** 
                  1. Nel Tab 6 (Entrate), hai digitato i dati ma **non hai cliccato il bottone rosso di salvataggio** (verifica la presenza di eventuali triangolini rossi o scritte `None` nelle celle).
                  2. Le uscite non appaiono perché non sono state registrate nella sezione *Uscite* del Tab 6 con una data e una colonna di importo valide (`Importo_Netto`), ma sono state lasciate come semplici stime di budget.
                * **Come risolverlo:** Clicca sempre i bottoni di salvataggio dedicati e compila le date puntuali nel registro contabile.
                """)

            with st.expander("👷 3. Allarmi di Sovraccarico Risorse nel Radar"):
                st.markdown("""
                * **La Situazione:** Il Radar di sinistra urla che l'Impresa X è in sovraccarico.
                * **Perché accade:** Due lavorazioni distinte assegnate alla stessa risorsa hanno date d'inizio e fine sovrapposte nel Tab 1.
                * **La Soluzione:** Se la sovrapposizione è voluta (es. l'impresa ha più squadre che lavorano in contemporanea su stanze diverse), non devi modificare il calendario: ti basta cliccare il bottone **'👁️ Consenti Sovrapposizione'** sotto l'allarme per registrarla come eccezione autorizzata.
                """)
                
            with st.expander("⏳ 4. Il percorso critico (CPM) non evidenzia i ritardi corretti"):
                st.markdown("""
                * **La Situazione:** Un'attività subisce un ritardo notevole, ma il Tab 3 (Grafo e Matrice) non la contrassegna come parte del percorso critico o non sposta la data di fine progetto.
                * **Perché accade:** Le dipendenze (*Predecessori*) non sono state collegate correttamente tramite gli ID WBS (es. manca il legame logico di fine-inizio tra le lavorazioni propedeutiche). Senza la catena dei predecessori, il motore considera le attività slegate e parallele.
                * **La Soluzione:** Verifica nel Tab 1 che ogni attività successiva abbia il proprio ID predecessore correttamente indicato (es. `2.1` come propedeutica alla `2.2`), in modo da permettere all'algoritmo CPM di ricalcolare la catena critica.
                """)

            with st.expander("📊 5. Valore Pianificato (PV) a zero nonostante le date inserite"):
                st.markdown("""
                * **La Situazione:** Il cruscotto dell'Earned Value (Tab 5) mostra un PV (Planned Value) pari a zero o incoerente con la data odierna di controllo.
                * **Perché accade:** Il budget (`BAC_Budget`) è stato assegnato solo alle attività di sintesi (i capitoli padre) anziché alle singole attività "foglia" esecutive, oppure le date di inizio e fine previste non rientrano nell'intervallo temporale analizzato.
                * **La Soluzione:** Assegna sempre i budget di spesa esclusivamente ai nodi foglia dell'albero WBS e assicurati che il cronoprogramma copra correttamente la linea del tempo corrente.
                """)

            with st.expander("🔄 6. Modifiche strutturali dell'albero e perdita dei dati associati"):
                st.markdown("""
                * **La Situazione:** Spostando, eliminando o riorganizzando i rami della WBS, alcune registrazioni contabili o assegnazioni di risorse sembrano svanire o non puntano più alla voce corretta.
                * **Perché accade:** Poiché il database relaziona costi (Tab 6), rischi (Tab 8) e CAPA (Tab 7) agli ID WBS specifici, la cancellazione o la modifica drastica di un codice ID interrompe la chiave di collegamento esterna.
                * **La Soluzione:** Prima di procedere con ristrutturazioni profonde dell'albero WBS (tramite la funzione di rinumerazione), esporta sempre un backup preventivo del file JSON del progetto per sicurezza.
                """)

            with st.expander("📉 7. L'indice SPI o CPI mostra valori anomali (> 2.0 o NaN)"):
                st.markdown("""
                * **La Situazione:** Nel cruscotto EVM (Tab 5) gli indicatori di performance temporale (SPI) o economica (CPI) assumono valori assurdi, pari a zero, o restituiscono un errore di calcolo.
                * **Perché accade:** L'Earned Value (EV) si calcola moltiplicando il Budget (`BAC`) per la percentuale di completamento. Se un'attività ha un budget pari a zero (`BAC = 0`) ma un avanzamento del 100%, o viceversa, la divisione matematica va in crisi per assenza di denominatore.
                * **La Soluzione:** Assicurati che **tutte** le attività abbiano un budget (`BAC_Budget`) maggiore di zero prima di iniziare a certificarne l'avanzamento fisico.
                """)

            with st.expander("🗂️ 8. Duplicazione o sfasamento dei file JSON caricati"):
                st.markdown("""
                * **La Situazione:** Caricando un vecchio file di progetto salvato in formato `.json`, noti che alcune tabelle (come OBS o i Rischi) si azzerano o non corrispondono più.
                * **Perché accade:** Stai tentando di caricare un file JSON generato da una versione precedente dell'applicazione che non conteneva ancora le chiavi dei nuovi database (es. i campi `sal` o `conflitti_ignorati`).
                * **La Soluzione:** Quando aggiorni l'applicazione con nuove sezioni, ricordati di esportare un nuovo file di progetto master "pulito" in modo che includa la struttura dati aggiornata di tutti i tab.
                """)

            with st.expander("⚠️ 9. Il Fondo Imprevisti del Tab 8 non copre i rischi attivi"):
                st.markdown("""
                * **La Situazione:** Hai inserito diversi rischi nella Matrice (Tab 8) con punteggi di impatto elevati, ma il budget di riserva calcolato non varia o sembra disallineato.
                * **Perché accade:** Il calcolo del fondo imprevisti si basa sullo stato di mitigazione e sul valore economico associato ai singoli rischi. Se i campi d'importo dei rischi sono stati lasciati a zero, il motore di rischio li considera solo come eventi qualitativi senza impatto di cassa.
                * **La Soluzione:** Compila sempre la stima economica dell'impatto all'interno della scheda di rischio nel Tab 8 per permettere al sistema di dimensionare correttamente le risorse di riserva.
                """)

            with st.expander("🚨 10. Allarme Critico Combinato nel Radar (Intervento RUP Richiesto)"):
                st.markdown("""
                * **La Situazione:** Nel cruscotto del Radar compare un vistoso messaggio di errore rosso che richiede l'intervento immediato del RUP, indicando che una specifica lavorazione è in ritardo e citando testualmente un rischio associato.
                * **Perché accade:** Non è un errore del programma, ma il **motore di controllo incrociato (Gianfry)** in azione. L'app ha scansionato l'intero progetto e ha rilevato una "tempesta perfetta": 
                  1. Un'attività nel Tab 1 ha superato la sua data di *Fine Prevista* ma non è ancora certificata al 100% (ritardo cronico sul campo).
                  2. Questa stessa attività è collegata a un evento nella Matrice dei Rischi (Tab 8) che è classificato con impatto *Alto* o *Critico* e risulta ancora nello stato *Aperto*.
                * **La Soluzione:** Il sistema ti sta avvisando che un rischio grave si sta materializzando a causa di un ritardo esecutivo. Per far rientrare l'allarme, il Direttore Lavori o il RUP devono adottare le misure di mitigazione previste; dopodiché basterà andare nel Tab 8 e commutare lo stato di quel rischio su **'Chiuso'**, oppure aggiornare l'effettivo completamento della WBS al 100% nel Tab 1.
                """)

            with st.expander("🏛️ 11. La data di 'Inizio Effettivo' scompare dopo il salvataggio (Blocco Amministrativo)"):
                st.markdown("""
                * **La Situazione:** Inserisci una data nella colonna *Inizio Effettivo* del Tab 1, clicchi "Salva", ma la cella torna immediatamente vuota e compare un banner di errore rosso.
                * **Perché accade:** L'attività in questione è soggetta a un **Cancello Amministrativo** (es. *Deposito Genio Civile* o *Autorizzazione Paesaggistica*). Il software, conformemente al Codice degli Appalti, impedisce la registrazione formale dell'inizio delle lavorazioni se manca il nulla osta istituzionale.
                * **La Soluzione:** L'inizio dei lavori può essere sbloccato solo da un utente autorizzato (RUP o DL). Quando l'ente preposto rilascia il protocollo ufficiale, entra nel Tab 1 e metti la spunta sulla colonna **'✅ Vincolo Assolto'**. Al salvataggio successivo, il sistema ti permetterà di inserire liberamente la data di Inizio Effettivo.
                """)

            with st.expander("🚨 12. L'Allerta Finanziaria dopo l'approvazione di una variante"):
                st.markdown("""
                * **La Situazione:** Nel Tab 9 (Audit Trail) il RUP ha approvato un Ticket di tipo "Richiesta di Variante", ma spostandoti nel Tab 1 (WBS) o nel Tab 5 (EVM) compare un avviso rosso fisso di "Allerta Finanziaria".
                * **Perché accade:** Hai formalmente approvato la variante, ma hai lasciato vuoti i campi "Variazione Costi (€)" e/o "Variazione Tempi (gg)". Il motore EVM sa che la contabilità contrattuale del cantiere è cambiata, ma non ha i numeri esatti per ricalcolare il Budget (BAC) e il cronoprogramma. È un "cane da guardia" per evitare che vengano concessi extra-budget a parole senza aggiornare la matematica del progetto.
                * **La Soluzione:** Torna nel Tab 9, compila gli importi esatti in Euro e i giorni concessi, poi clicca "Registra Risposte RUP". Il sistema inietterà automaticamente quei valori nel Budget della WBS interessata (nel Tab 1), sposterà in avanti la data di *Fine Prevista*, ricalcolerà tutti gli indicatori nel Tab 5 e spegnerà l'allarme in modo definitivo. 
                *(Nota tecnica: il sistema possiede una memoria interna, quindi i soldi non verranno mai sommati due volte, nemmeno se continui a premere Salva).*
                """)

            with st.expander("🚦 13. Perché il Gantt o i KPI segnano 'Arancione' se l'SPI è 0.99?"):
                st.markdown("""
                * **La Situazione:** Il tuo indice di efficienza temporale (SPI) è a 0.99. Visivamente e logicamente verrebbe da dire: *"È praticamente 1, dammi il verde!"*, ma il cruscotto e la barra del Gantt si colorano inesorabilmente di Arancione.
                * **Perché accade:** Entra in gioco la spietata rigidità matematica del motore EVM. Le regole assolute del software sono: 🟢 **Verde** per SPI >= 1.00, 🟠 **Arancione** per SPI compreso tra 0.90 e 0.99, e 🔴 **Rosso** sotto lo 0.90. Poiché 0.99 è strettamente minore di 1.00, la macchina, senza un briciolo di flessibilità umana, lo declassa a "Lieve Ritardo". Ti sta matematicamente dicendo: *"Sei in ritardo dell'1% rispetto al piano teorico"*.
                * **La Soluzione:** Nel Project Management "reale", spaccare il secondo al 100% è pressoché impossibile. Di solito si adotta una "soglia di tolleranza" operativa. L'app nasce con la massima severità di default per non nascondere alcuna deviazione, ma se questa pignoleria algoritmica dovesse risultare eccessiva, i parametri del software possono essere ricalibrati in qualsiasi momento (modificando le funzioni del Gantt e dei KPI) per considerare "Verde" tutto ciò che si mantiene sopra una soglia di tolleranza dello 0.95.
                """)

        # --- SEZIONE 5: ROADMAP VERSIONE 2.0 ---
        with t_sec5:
            st.subheader("🚀 Prossimi Sviluppi (Roadmap v2.0)")
            st.markdown("Ecco le funzioni di livello Enterprise che verranno integrate nelle prossime release autonome:")
            
            st.info("""
            * **📥 Importatore Nativo PriMus (.csv / .xls):** Integrazione diretta per mappare i computi metrici estimativi di ACCA software direttamente sull'albero WBS e sui budget di cantiere.
            * **🌐 Viewer BIM 4D / SLAM integrato:** Collegamento delle nuvole di punti 3D e dei virtual tour laser scanner direttamente alle singole voci di stato d'avanzamento.
            * **📄 Generatore Automatico del Giornale dei Lavori:** Esportazione in PDF impaginata con loghi, verbali di cantiere, firme e grafici di cash flow pronti per la DL.
            * **📊 Modulo Finanziario Immobiliare Avanzato:** Calcolo automatico del VAN (Valore Attivo Netto) e del TIR (Tasso Interno di Rendimento) per operazioni di sviluppo immobiliare.
            """)
